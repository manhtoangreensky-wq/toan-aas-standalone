"""Bounded, private Web-native document operations.

The first operations are PDF Split, PDF Merge, lossless PDF Optimize and
Image to PDF. They reuse only verified, owner-scoped Asset Vault inputs, copy
them into an isolated operation staging area, and create separately stored
output attachments. They never call the Bot, a provider, wallet, PayOS, or a
browser-supplied path.
"""

from __future__ import annotations

from contextlib import ExitStack
from datetime import datetime, timedelta, timezone
import hashlib
import hmac
from io import BytesIO
import json
import math
import os
from pathlib import Path
import re
import threading
import uuid
from typing import Any
import warnings
from zipfile import BadZipFile, ZIP_DEFLATED, ZipFile, ZipInfo

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel, ConfigDict, Field, StrictStr, field_validator
from starlette.concurrency import run_in_threadpool

from copyfast_auth import _record_audit, _request_id, envelope, require_account, require_csrf
from copyfast_image_runtime import image_decoder_capacity
from copyfast_db import (
    asset_vault_directory,
    asset_vault_enabled,
    document_operations_directory,
    document_operations_enabled,
    image_ocr_enabled,
    ensure_copyfast_schema,
    image_to_pdf_enabled,
    pdf_to_images_enabled,
    pdf_to_word_enabled,
    transaction,
    utc_now,
)


router = APIRouter(prefix="/api/v1/document-operations", tags=["Web Document Operations"])

PDF_SPLIT_KIND = "pdf_split"
PDF_MERGE_KIND = "pdf_merge"
PDF_OPTIMIZE_KIND = "pdf_optimize"
IMAGE_TO_PDF_KIND = "image_to_pdf"
PDF_TO_IMAGES_KIND = "pdf_to_images"
PDF_TO_WORD_KIND = "pdf_to_word_text"
IMAGE_OCR_KIND = "image_ocr"
SUPPORTED_KINDS = frozenset({
    PDF_SPLIT_KIND,
    PDF_MERGE_KIND,
    PDF_OPTIMIZE_KIND,
    IMAGE_TO_PDF_KIND,
    PDF_TO_IMAGES_KIND,
    PDF_TO_WORD_KIND,
    IMAGE_OCR_KIND,
})
OPERATION_STATES = frozenset({"queued", "processing", "completed", "failed", "unavailable", "guarded"})
IDEMPOTENCY_PATTERN = re.compile(r"^[A-Za-z0-9._:-]{12,160}$")
UUID_PATTERN = re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[1-5][0-9a-f]{3}-[89ab][0-9a-f]{3}-[0-9a-f]{12}$", re.IGNORECASE)
PAGE_RANGE_PATTERN = re.compile(r"^\s*(\d+)\s*(?:-\s*(\d+)\s*)?$")
ASSET_STORAGE_KEY_PATTERN = re.compile(r"^objects/[0-9a-f]{32}\.blob$")
OUTPUT_STORAGE_KEY_PATTERN = re.compile(r"^outputs/[0-9a-f]{32}\.(?P<suffix>pdf|docx|png|txt|zip)$")
CHUNK_BYTES = 1024 * 1024
MAX_INPUT_BYTES = 20 * 1024 * 1024  # Mirrors the current Bot PDF limit.
MAX_PAGES = 30  # Mirrors Bot `DOC_MAX_PAGES` and bounds parser work.
MAX_MERGE_SOURCES = 8  # Web safety bound; Bot itself only requires at least two.
MAX_MERGE_INPUT_BYTES = 40 * 1024 * 1024  # Aggregate bound beyond per-file 20 MiB.
# Image-to-PDF deliberately keeps a lower bounded batch than the legacy Bot.
# Each decoded source can use substantially more memory than its compressed
# upload size, so a 1–8 page Web batch is both useful and predictable on the
# shared application runtime.  The Bot's wider 20-file guided flow is recorded
# in the migration inventory rather than copied without its missing guards.
MAX_IMAGE_PDF_SOURCES = 8
MAX_IMAGE_PDF_INPUT_BYTES = 40 * 1024 * 1024
# A decoded raster uses far more memory than its compressed upload.  Keep the
# per-page maximum comfortably below a 512 MiB service, then serialize this
# decoder-backed operation per process below.  This is intentionally stricter
# than the old Bot helper, whose runtime and operational limits are different.
MAX_IMAGE_PIXELS_PER_SOURCE = 16 * 1024 * 1024
MAX_IMAGE_PIXELS_TOTAL = 32 * 1024 * 1024
MAX_IMAGE_DIMENSION = 7_680
MAX_IMAGE_ASPECT_RATIO = 12
IMAGE_TO_PDF_MAX_CONCURRENT = 1
# Keep this legacy-named alias for the documented Image → PDF contract and
# tests, but back it with the process-wide Pillow gate shared with Resize &
# Aspect Studio. Separate per-feature semaphores would allow two 16 MP decodes
# to run together and defeat the memory boundary.
_IMAGE_TO_PDF_CAPACITY = image_decoder_capacity()
PDF_TO_WORD_MAX_CONCURRENT = 1
_PDF_TO_WORD_CAPACITY = threading.BoundedSemaphore(value=PDF_TO_WORD_MAX_CONCURRENT)
# PDF rendering has a separate native decoder/runtime from pypdf text or
# structure work. One render at a time keeps peak bitmap memory predictable
# per Web process and cannot be bypassed by concurrent accounts.
PDF_TO_IMAGES_MAX_CONCURRENT = 1
_PDF_TO_IMAGES_CAPACITY = threading.BoundedSemaphore(value=PDF_TO_IMAGES_MAX_CONCURRENT)
# OCR opens the same bounded image raster as Image → PDF / Image Operations.
# Reuse the process-wide decoder gate rather than creating a second OCR-only
# semaphore that could allow two 16 MP inputs to decode at once.
IMAGE_OCR_MAX_CONCURRENT = 1
_IMAGE_OCR_CAPACITY = image_decoder_capacity()
IMAGE_OCR_LANGUAGES = frozenset({"auto", "vi", "en"})
MAX_IMAGE_OCR_CHARACTERS = 500_000
MAX_IMAGE_OCR_OUTPUT_BYTES = 2 * 1024 * 1024
PDF_TO_IMAGES_RENDER_SCALE = 2.0  # Mirrors Bot `/pdf_to_images` 2× render.
MAX_PDF_TO_IMAGES_PAGE_PIXELS = 8 * 1024 * 1024
MAX_PDF_TO_IMAGES_TOTAL_PIXELS = 48 * 1024 * 1024
MAX_PDF_TO_IMAGES_DIMENSION = 8_192
MAX_PDF_TO_IMAGES_PAGE_PNG_BYTES = 8 * 1024 * 1024
MAX_PDF_TO_IMAGES_TOTAL_PNG_BYTES = 32 * 1024 * 1024
MAX_PDF_TO_IMAGES_ZIP_MEMBERS = MAX_PAGES
# Keep archive metadata stable as well as the page order/names.  This makes a
# replayable Web artifact independently reproducible from the same rendered
# PNG bytes instead of leaking host clock or filesystem timestamps.
PDF_TO_IMAGES_ZIP_DATE_TIME = (1980, 1, 1, 0, 0, 0)
# Text extraction is bounded independently of PDF byte/page limits.  Text can
# expand disproportionately when the input contains repeated or malformed
# glyph streams.  The DOCX is a logical-text export, not an OCR/layout engine.
MAX_PDF_TO_WORD_CHARACTERS = 250_000
MAX_PDF_TO_WORD_PAGE_CHARACTERS = 25_000
MAX_PDF_TO_WORD_PARAGRAPHS = 10_000
MAX_DOCX_ARCHIVE_MEMBERS = 200
MAX_DOCX_UNCOMPRESSED_BYTES = 8 * 1024 * 1024
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
IMAGE_INPUT_MIME_BY_EXTENSION = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
}
# A result that is technically one byte smaller is not a professional
# optimization. Keep the original unless the verified final artifact saves a
# meaningful amount without a lossy engine or an external command runner.
MIN_OPTIMIZATION_SAVED_BYTES = 1024
MIN_OPTIMIZATION_SAVED_RATIO = 0.01
ORPHAN_RETENTION_SECONDS = 60 * 60
PDF_EXCLUDED_PAGE_KEYS = ("/Annots", "/AA", "/Metadata", "/PieceInfo", "/StructParents")

OUTPUT_SPEC_BY_KIND = {
    PDF_SPLIT_KIND: (".pdf", "application/pdf", "toan-aas-pdf-split.pdf"),
    PDF_MERGE_KIND: (".pdf", "application/pdf", "toan-aas-pdf-merged.pdf"),
    PDF_OPTIMIZE_KIND: (".pdf", "application/pdf", "toan-aas-pdf-optimized.pdf"),
    IMAGE_TO_PDF_KIND: (".pdf", "application/pdf", "toan-aas-images.pdf"),
    # The single-page PNG case is resolved dynamically in `_output_spec`;
    # multi-page outputs use this canonical ZIP contract.
    PDF_TO_IMAGES_KIND: (".zip", "application/zip", "toan-aas-pdf-pages.zip"),
    PDF_TO_WORD_KIND: (".docx", DOCX_MEDIA_TYPE, "toan-aas-pdf-text.docx"),
    IMAGE_OCR_KIND: (".txt", "text/plain; charset=utf-8", "toan-aas-image-ocr.txt"),
}

OPERATION_SELECT = """id, source_asset_id, project_id, kind, state, requested_page_range,
                      selected_start_page, selected_end_page, source_page_count, output_page_count,
                      original_filename, content_type, byte_size, created_at, queued_at, started_at,
                      completed_at, updated_at, failure_code, storage_key, sha256, source_sha256,
                      source_byte_size, source_count"""


class DocumentOperationError(Exception):
    """A known safe failure that must never expose a parser/infrastructure trace."""

    def __init__(self, message: str, *, code: str = "DOCUMENT_OPERATION_INVALID"):
        super().__init__(message)
        self.public_message = message
        self.code = code


class PdfSplitRequest(BaseModel):
    source_asset_id: str = Field(min_length=36, max_length=36)
    page_range: str = Field(min_length=1, max_length=32)
    idempotency_key: str = Field(min_length=12, max_length=160)

    @field_validator("source_asset_id")
    @classmethod
    def valid_source_asset_id(cls, value: str) -> str:
        return _uuid(value, label="Asset Vault ID")

    @field_validator("page_range")
    @classmethod
    def valid_page_range(cls, value: str) -> str:
        compact = str(value or "").strip()
        match = PAGE_RANGE_PATTERN.fullmatch(compact)
        if not match:
            raise ValueError("Khoảng trang phải là một trang hoặc dải liên tiếp, ví dụ 2 hoặc 2-5")
        start = int(match.group(1))
        end = int(match.group(2) or start)
        # Keep Bot's useful reverse-range behavior, while canonicalizing the
        # request before it enters the idempotency fingerprint/database.
        # `5-2` and `2-5` are one user intent, not two output artifacts.
        if start > end:
            start, end = end, start
        return str(start) if start == end else f"{start}-{end}"

    @field_validator("idempotency_key")
    @classmethod
    def valid_idempotency_key(cls, value: str) -> str:
        return _idempotency_key(value)


class PdfMergeRequest(BaseModel):
    source_asset_ids: list[str] = Field(min_length=2, max_length=MAX_MERGE_SOURCES)
    idempotency_key: str = Field(min_length=12, max_length=160)

    @field_validator("source_asset_ids")
    @classmethod
    def valid_source_asset_ids(cls, values: list[str]) -> list[str]:
        normalized = [_uuid(value, label="Asset Vault ID") for value in values]
        if len(normalized) < 2:
            raise ValueError("Cần ít nhất hai PDF private để gộp")
        if len(set(normalized)) != len(normalized):
            raise ValueError("Mỗi PDF nguồn chỉ được chọn một lần trong cùng thao tác gộp")
        return normalized

    @field_validator("idempotency_key")
    @classmethod
    def valid_idempotency_key(cls, value: str) -> str:
        return _idempotency_key(value)


class PdfOptimizeRequest(BaseModel):
    source_asset_id: str = Field(min_length=36, max_length=36)
    idempotency_key: str = Field(min_length=12, max_length=160)

    @field_validator("source_asset_id")
    @classmethod
    def valid_source_asset_id(cls, value: str) -> str:
        return _uuid(value, label="Asset Vault ID")

    @field_validator("idempotency_key")
    @classmethod
    def valid_idempotency_key(cls, value: str) -> str:
        return _idempotency_key(value)


class PdfToWordRequest(BaseModel):
    """One owner-scoped text-bearing PDF becomes a fresh private DOCX."""

    source_asset_id: str = Field(min_length=36, max_length=36)
    idempotency_key: str = Field(min_length=12, max_length=160)

    @field_validator("source_asset_id")
    @classmethod
    def valid_source_asset_id(cls, value: str) -> str:
        return _uuid(value, label="Asset Vault ID")

    @field_validator("idempotency_key")
    @classmethod
    def valid_idempotency_key(cls, value: str) -> str:
        return _idempotency_key(value)


class PdfToImagesRequest(BaseModel):
    """One verified private PDF becomes one PNG or a deterministic PNG ZIP."""

    source_asset_id: str = Field(min_length=36, max_length=36)
    idempotency_key: str = Field(min_length=12, max_length=160)

    @field_validator("source_asset_id")
    @classmethod
    def valid_source_asset_id(cls, value: str) -> str:
        return _uuid(value, label="Asset Vault ID")

    @field_validator("idempotency_key")
    @classmethod
    def valid_idempotency_key(cls, value: str) -> str:
        return _idempotency_key(value)


class ImageToPdfRequest(BaseModel):
    """An explicit private image order; every source becomes one PDF page."""

    source_asset_ids: list[str] = Field(min_length=1, max_length=MAX_IMAGE_PDF_SOURCES)
    idempotency_key: str = Field(min_length=12, max_length=160)

    @field_validator("source_asset_ids")
    @classmethod
    def valid_source_asset_ids(cls, values: list[str]) -> list[str]:
        normalized = [_uuid(value, label="Asset Vault ID") for value in values]
        if not normalized:
            raise ValueError("Cần ít nhất một ảnh private để tạo PDF")
        if len(set(normalized)) != len(normalized):
            raise ValueError("Mỗi ảnh nguồn chỉ được chọn một lần trong cùng thao tác")
        return normalized

    @field_validator("idempotency_key")
    @classmethod
    def valid_idempotency_key(cls, value: str) -> str:
        return _idempotency_key(value)


class ImageOcrRequest(BaseModel):
    """One owner-scoped image becomes one private, verified text artifact.

    The browser cannot provide an idempotency key, URL, path, bytes or
    provider option.  The server derives an idempotency identity from the
    immutable Asset Vault revision and requested local language mode.
    """

    model_config = ConfigDict(extra="forbid", str_strip_whitespace=True)

    source_asset_id: StrictStr = Field(min_length=36, max_length=36)
    language: StrictStr = Field(min_length=2, max_length=4)

    @field_validator("source_asset_id")
    @classmethod
    def valid_source_asset_id(cls, value: StrictStr) -> str:
        return _uuid(value, label="Asset Vault ID")

    @field_validator("language")
    @classmethod
    def valid_language(cls, value: StrictStr) -> str:
        raw = str(value or "").strip()
        normalized = raw.lower()
        if raw != normalized or normalized not in IMAGE_OCR_LANGUAGES:
            raise ValueError("Ngôn ngữ OCR không hợp lệ")
        return normalized


def _require_enabled() -> None:
    if not document_operations_enabled() or not asset_vault_enabled():
        raise HTTPException(
            status_code=503,
            detail="Document Operations cần Asset Vault private và storage đầu ra riêng đã được bật",
        )


def _require_image_to_pdf_enabled() -> None:
    """Keep the decoder-backed operation behind its own explicit circuit breaker."""
    _require_enabled()
    if not image_to_pdf_enabled():
        raise HTTPException(
            status_code=503,
            detail="Ảnh → PDF chưa được bật; cần private storage và WEBAPP_IMAGE_TO_PDF_ENABLED",
        )


def _require_image_ocr_enabled() -> None:
    """Keep local OCR explicitly opt-in and independent from PDF utilities."""

    _require_enabled()
    if not image_ocr_enabled():
        raise HTTPException(
            status_code=503,
            detail="OCR ảnh private chưa được bật; cần WEBAPP_DOCUMENT_OCR_IMAGE_ENABLED và local Tesseract runtime",
        )


def _require_pdf_to_word_enabled() -> None:
    """Keep the DOCX writer behind an independent, fail-closed feature gate."""
    _require_enabled()
    if not pdf_to_word_enabled():
        raise HTTPException(
            status_code=503,
            detail="PDF có text → Word chưa được bật; cần private storage và WEBAPP_PDF_TO_WORD_ENABLED",
        )


def _require_pdf_to_images_enabled() -> None:
    """Keep PDFium raster rendering behind its own explicit circuit breaker."""
    _require_enabled()
    if not pdf_to_images_enabled():
        raise HTTPException(
            status_code=503,
            detail="PDF → ảnh chưa được bật; cần private storage và WEBAPP_PDF_TO_IMAGES_ENABLED",
        )


def _reserve_image_to_pdf_capacity() -> None:
    """Allow only one decoder-heavy image batch per Web process.

    The HTTP/IP limiter protects the request surface, but it cannot bound
    aggregate decoded memory across distinct clients.  This small process-wide
    semaphore is deliberately fail-fast: a second batch receives an honest
    retry response before it can create a database row or staging artifact.
    """
    if not _IMAGE_TO_PDF_CAPACITY.acquire(blocking=False):
        raise HTTPException(
            status_code=429,
            detail="Ảnh → PDF đang bận xử lý một lô khác; vui lòng thử lại sau ít phút",
        )


def _reserve_image_ocr_capacity() -> None:
    """Reserve the shared decoded-image slot before creating an OCR row."""

    if not _IMAGE_OCR_CAPACITY.acquire(blocking=False):
        raise HTTPException(
            status_code=429,
            detail="OCR ảnh đang bận xử lý một ảnh khác; vui lòng thử lại sau ít phút",
        )


def _reserve_pdf_to_word_capacity() -> None:
    """Serialize bounded text extraction/DOCX generation per Web process."""
    if not _PDF_TO_WORD_CAPACITY.acquire(blocking=False):
        raise HTTPException(
            status_code=429,
            detail="PDF có text → Word đang bận xử lý một tệp khác; vui lòng thử lại sau ít phút",
        )


def _reserve_pdf_to_images_capacity() -> None:
    """Serialize bounded PDF bitmap rendering per Web process."""
    if not _PDF_TO_IMAGES_CAPACITY.acquire(blocking=False):
        raise HTTPException(
            status_code=429,
            detail="PDF → ảnh đang bận xử lý một tệp khác; vui lòng thử lại sau ít phút",
        )


def ensure_document_operations_runtime() -> None:
    """Fail closed only for enabled parsers, leaving disabled surfaces inert."""
    if not document_operations_enabled():
        return
    try:
        from pypdf import PdfReader, PdfWriter  # noqa: F401
    except ImportError as exc:  # pragma: no cover - deployment configuration
        raise RuntimeError("Document Operations cần dependency pypdf") from exc
    if image_to_pdf_enabled():
        _image_classes()
    if pdf_to_word_enabled():
        _word_classes()
    if pdf_to_images_enabled():
        _pdf_to_images_classes()


def _pdf_classes():
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:  # pragma: no cover - ensured at startup when enabled
        raise DocumentOperationError("Document Operations chưa có runtime PDF an toàn", code="PDF_RUNTIME_UNAVAILABLE") from exc
    return PdfReader, PdfWriter


def _image_classes():
    """Load the decoder lazily so a disabled feature cannot alter startup."""
    try:
        from PIL import Image, ImageFile, ImageOps, UnidentifiedImageError
    except ImportError as exc:  # pragma: no cover - deployment configuration
        raise DocumentOperationError(
            "Ảnh → PDF chưa có runtime Pillow an toàn",
            code="IMAGE_RUNTIME_UNAVAILABLE",
        ) from exc
    return Image, ImageFile, ImageOps, UnidentifiedImageError


def _image_ocr_runtime() -> tuple[Any, frozenset[str]]:
    """Load an explicitly local OCR runtime without making startup depend on it.

    `pytesseract` is only a local adapter around a service-installed Tesseract
    binary.  A missing binary or language directory is a normal readiness
    state for this optional feature, not a reason to prevent the entire Web
    App from starting.  The request handler turns this known condition into a
    guarded, non-delivery response.
    """

    try:
        import pytesseract
    except ImportError as exc:
        raise DocumentOperationError(
            "OCR ảnh chưa có local Tesseract runtime đã được cấu hình.",
            code="OCR_RUNTIME_UNAVAILABLE",
        ) from exc
    try:
        # Make a small explicit readiness call before any private image is
        # decoded. `get_languages` also proves the configured tessdata path is
        # usable without exposing it to the browser.
        pytesseract.get_tesseract_version()
        languages = frozenset(
            str(item or "").strip().lower()
            for item in pytesseract.get_languages(config="")
            if str(item or "").strip()
        )
    except Exception as exc:
        raise DocumentOperationError(
            "OCR ảnh chưa có local Tesseract runtime sẵn sàng.",
            code="OCR_RUNTIME_UNAVAILABLE",
        ) from exc
    if not languages:
        raise DocumentOperationError(
            "OCR ảnh chưa có language pack local sẵn sàng.",
            code="OCR_RUNTIME_UNAVAILABLE",
        )
    return pytesseract, languages


def _image_ocr_language(language: str, languages: frozenset[str]) -> str:
    """Resolve only documented public language choices to installed packs."""

    available = frozenset(str(item or "").strip().lower() for item in languages if str(item or "").strip())
    if language == "vi":
        if "vie" not in available:
            raise DocumentOperationError(
                "OCR ảnh chưa có language pack tiếng Việt local.",
                code="OCR_LANGUAGE_UNAVAILABLE",
            )
        return "vie"
    if language == "en":
        if "eng" not in available:
            raise DocumentOperationError(
                "OCR ảnh chưa có language pack tiếng Anh local.",
                code="OCR_LANGUAGE_UNAVAILABLE",
            )
        return "eng"
    if "vie" in available and "eng" in available:
        return "vie+eng"
    if "eng" in available:
        return "eng"
    if "vie" in available:
        return "vie"
    raise DocumentOperationError(
        "OCR ảnh chưa có language pack local phù hợp.",
        code="OCR_LANGUAGE_UNAVAILABLE",
    )


def _word_classes():
    """Load python-docx only for the explicitly enabled DOCX exporter."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - deployment configuration
        raise DocumentOperationError(
            "PDF có text → Word chưa có runtime DOCX an toàn",
            code="PDF_TO_WORD_RUNTIME_UNAVAILABLE",
        ) from exc
    return Document


def _pdf_to_images_classes():
    """Load the BSD/Apache PDFium renderer only for the enabled feature.

    The Bot historically uses PyMuPDF.  The standalone Web App deliberately
    uses PDFium through pypdfium2 so the Web package has a separately reviewed
    raster dependency while retaining the same 2× public result contract.
    """
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:  # pragma: no cover - deployment configuration
        raise DocumentOperationError(
            "PDF → ảnh chưa có runtime PDFium an toàn",
            code="PDF_TO_IMAGES_RUNTIME_UNAVAILABLE",
        ) from exc
    return pdfium


def _maximum_output_bytes() -> int:
    raw = os.environ.get("WEBAPP_DOCUMENT_OPERATIONS_MAX_OUTPUT_MB", "20").strip()
    try:
        megabytes = int(raw)
    except ValueError:
        megabytes = 20
    return max(1, min(megabytes, 50)) * 1024 * 1024


def _maximum_account_bytes() -> int:
    raw = os.environ.get("WEBAPP_DOCUMENT_OPERATIONS_QUOTA_MB", "100").strip()
    try:
        megabytes = int(raw)
    except ValueError:
        megabytes = 100
    return max(1, min(megabytes, 5_000)) * 1024 * 1024


def _uuid(value: str, *, label: str) -> str:
    candidate = str(value or "").strip()
    if not UUID_PATTERN.fullmatch(candidate):
        raise HTTPException(status_code=422, detail=f"{label} không hợp lệ")
    return str(uuid.UUID(candidate))


def _idempotency_key(value: str | None) -> str:
    key = str(value or "").strip()
    if not IDEMPOTENCY_PATTERN.fullmatch(key):
        raise HTTPException(status_code=422, detail="Idempotency key không hợp lệ")
    return key


def _asset_path(root: Path, storage_key: str) -> Path:
    if not ASSET_STORAGE_KEY_PATTERN.fullmatch(storage_key):
        raise RuntimeError("Storage key Asset Vault không hợp lệ")
    candidate = (root / storage_key).resolve()
    try:
        candidate.relative_to(root.resolve())
    except ValueError as exc:
        raise RuntimeError("Storage key Asset Vault vượt ngoài thư mục riêng") from exc
    return candidate


def _output_path(root: Path, storage_key: str, *, expected_suffix: str | None = None) -> Path:
    match = OUTPUT_STORAGE_KEY_PATTERN.fullmatch(storage_key)
    if not match or (expected_suffix is not None and f".{match.group('suffix')}" != expected_suffix):
        raise RuntimeError("Storage key Document Operation không hợp lệ")
    candidate = (root / storage_key).resolve()
    try:
        candidate.relative_to(root.resolve())
    except ValueError as exc:
        raise RuntimeError("Storage key Document Operation vượt ngoài thư mục riêng") from exc
    return candidate


def _staging_path(root: Path, suffix: str) -> Path:
    staging = _private_operation_directory(root, ".staging")
    return staging / f"{uuid.uuid4().hex}{suffix}"


def _private_operation_directory(root: Path, name: str) -> Path:
    """Create a real child directory without following a later symlink swap."""
    if name not in {".staging", "outputs"}:
        raise RuntimeError("Thư mục Document Operation không hợp lệ")
    private_root = root.resolve()
    candidate = private_root / name
    if candidate.exists() and candidate.is_symlink():
        raise RuntimeError("Thư mục Document Operation không được là symbolic link")
    candidate.mkdir(parents=True, exist_ok=True)
    if not candidate.is_dir() or candidate.is_symlink():
        raise RuntimeError("Thư mục Document Operation không hợp lệ")
    resolved = candidate.resolve()
    try:
        resolved.relative_to(private_root)
    except ValueError as exc:
        raise RuntimeError("Thư mục Document Operation vượt ngoài storage riêng") from exc
    return resolved


def _safe_unlink(path: Path | None) -> None:
    if path is None:
        return
    try:
        if path.is_file() and not path.is_symlink():
            path.unlink()
    except OSError:
        pass


def _verify_file(path: Path, *, expected_bytes: int, expected_digest: str) -> bool:
    try:
        if not path.is_file() or path.is_symlink() or path.stat().st_size != expected_bytes:
            return False
        digest = hashlib.sha256()
        with path.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
        return hmac.compare_digest(digest.hexdigest(), expected_digest)
    except OSError:
        return False


def _copy_verified_source(source: Path, destination: Path, *, expected_bytes: int, expected_digest: str) -> None:
    """Copy a verified Asset Vault blob so parser input cannot race its source."""
    total = 0
    digest = hashlib.sha256()
    prefix = b""
    try:
        if not source.is_file() or source.is_symlink() or source.stat().st_size != expected_bytes:
            raise DocumentOperationError("PDF nguồn không còn sẵn sàng", code="PDF_SOURCE_UNAVAILABLE")
        with source.open("rb") as read_stream, destination.open("xb") as write_stream:
            while True:
                chunk = read_stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                total += len(chunk)
                if total > MAX_INPUT_BYTES:
                    raise DocumentOperationError("PDF nguồn vượt giới hạn 20 MB", code="PDF_INPUT_TOO_LARGE")
                if len(prefix) < 8:
                    prefix += chunk[: 8 - len(prefix)]
                digest.update(chunk)
                write_stream.write(chunk)
    except OSError as exc:
        raise DocumentOperationError("Không thể đọc PDF nguồn riêng tư", code="PDF_SOURCE_UNAVAILABLE") from exc
    if total != expected_bytes or not hmac.compare_digest(digest.hexdigest(), expected_digest) or not prefix.startswith(b"%PDF-"):
        raise DocumentOperationError("PDF nguồn không vượt qua kiểm tra integrity", code="PDF_SOURCE_UNAVAILABLE")


def _image_magic_matches(extension: str, prefix: bytes) -> bool:
    """Repeat the Asset Vault signature check after an isolated copy."""
    if extension in {".jpg", ".jpeg"}:
        return prefix.startswith(b"\xff\xd8\xff")
    if extension == ".png":
        return prefix.startswith(b"\x89PNG\r\n\x1a\n")
    if extension == ".webp":
        return len(prefix) >= 12 and prefix.startswith(b"RIFF") and prefix[8:12] == b"WEBP"
    return False


def _copy_verified_image_source(
    source: Path,
    destination: Path,
    *,
    extension: str,
    expected_bytes: int,
    expected_digest: str,
) -> None:
    """Copy one image into private staging and verify bytes/hash/signature again."""
    total = 0
    digest = hashlib.sha256()
    prefix = b""
    try:
        if not source.is_file() or source.is_symlink() or source.stat().st_size != expected_bytes:
            raise DocumentOperationError("Ảnh nguồn không còn sẵn sàng", code="IMAGE_SOURCE_UNAVAILABLE")
        with source.open("rb") as read_stream, destination.open("xb") as write_stream:
            while True:
                chunk = read_stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                total += len(chunk)
                if total > MAX_INPUT_BYTES:
                    raise DocumentOperationError("Mỗi ảnh nguồn không được vượt quá 20 MB", code="IMAGE_INPUT_TOO_LARGE")
                if len(prefix) < 16:
                    prefix += chunk[: 16 - len(prefix)]
                digest.update(chunk)
                write_stream.write(chunk)
    except OSError as exc:
        raise DocumentOperationError("Không thể đọc ảnh nguồn riêng tư", code="IMAGE_SOURCE_UNAVAILABLE") from exc
    if (
        total != expected_bytes
        or not hmac.compare_digest(digest.hexdigest(), expected_digest)
        or not _image_magic_matches(extension, prefix)
    ):
        raise DocumentOperationError("Ảnh nguồn không vượt qua kiểm tra integrity", code="IMAGE_SOURCE_UNAVAILABLE")


def _selected_pages(page_range: str, page_count: int) -> tuple[list[int], int, int]:
    if page_count < 1 or page_count > MAX_PAGES:
        raise DocumentOperationError(
            f"PDF cần từ 1 đến {MAX_PAGES} trang để tách an toàn",
            code="PDF_PAGE_LIMIT",
        )
    match = PAGE_RANGE_PATTERN.fullmatch(str(page_range or "").strip())
    if not match:
        raise DocumentOperationError("Khoảng trang không hợp lệ", code="PDF_PAGE_RANGE")
    start = int(match.group(1))
    end = int(match.group(2) or start)
    if start > end:
        start, end = end, start
    if start < 1 or end > page_count:
        raise DocumentOperationError(f"Trang cần nằm trong phạm vi 1-{page_count}", code="PDF_PAGE_RANGE")
    selected = list(range(start - 1, end))
    if not selected or len(selected) > MAX_PAGES:
        raise DocumentOperationError("Số trang cần tách vượt giới hạn an toàn", code="PDF_PAGE_LIMIT")
    return selected, start, end


def _operation_public(row: tuple[Any, ...]) -> dict[str, Any]:
    state = str(row[4])
    kind = str(row[3])
    requested_page_range = str(row[5])
    # OCR stores its deliberately narrow language selector with operation
    # metadata. Expose only the allow-listed selector for an OCR history
    # item; never turn a request metadata field into a generic public input.
    ocr_language = None
    if kind == IMAGE_OCR_KIND and requested_page_range.startswith("ocr:"):
        candidate = requested_page_range.removeprefix("ocr:")
        if candidate in IMAGE_OCR_LANGUAGES:
            ocr_language = candidate
    byte_size = int(row[12]) if row[12] is not None else None
    source_byte_size = int(row[22]) if row[22] is not None else None
    source_count = max(1, int(row[23] or 1)) if len(row) > 23 else 1
    optimization_saved_bytes = (
        max(0, source_byte_size - byte_size)
        if kind == PDF_OPTIMIZE_KIND and source_byte_size is not None and byte_size is not None and state == "completed"
        else None
    )
    return {
        "id": str(row[0]),
        "source_asset_id": str(row[1]),
        "source_count": source_count,
        "project_id": str(row[2]) if row[2] else None,
        "kind": kind,
        "state": state,
        "requested_page_range": requested_page_range,
        "language": ocr_language,
        "selected_start_page": int(row[6]) if row[6] is not None else None,
        "selected_end_page": int(row[7]) if row[7] is not None else None,
        "source_page_count": int(row[8]) if row[8] is not None else None,
        "output_page_count": int(row[9]) if row[9] is not None else None,
        "original_filename": str(row[10]) if row[10] else None,
        "content_type": str(row[11]) if row[11] else None,
        "byte_size": byte_size,
        "created_at": str(row[13]),
        "queued_at": str(row[14]),
        "started_at": str(row[15]) if row[15] else None,
        "completed_at": str(row[16]) if row[16] else None,
        "updated_at": str(row[17]),
        "download_ready": state == "completed" and bool(row[19]) and byte_size is not None,
        # These derived values are exposed only for a completed optimize
        # artifact. They are safe account-owned measurements, not a storage
        # key, hash, path or source blob projection.
        "input_byte_size": source_byte_size if kind == PDF_OPTIMIZE_KIND and state == "completed" else None,
        "saved_bytes": optimization_saved_bytes,
        "saved_percent": (
            round((optimization_saved_bytes / source_byte_size) * 100, 1)
            if optimization_saved_bytes is not None and source_byte_size else None
        ),
        # Used only by _operation_response to choose a safe guarded message.
        # That helper strips private keys before sending the operation object
        # to a browser.
        "_failure_code": str(row[18] or ""),
    }


def _operation_not_found() -> dict[str, Any]:
    return envelope(
        False,
        "Không tìm thấy thao tác tài liệu thuộc Web account hiện tại.",
        status_name="guarded",
        error_code="WEB_DOCUMENT_OPERATION_NOT_FOUND",
    )


def _source_not_found() -> dict[str, Any]:
    return envelope(
        False,
        "Không tìm thấy PDF private đang hoạt động thuộc Web account hiện tại.",
        status_name="guarded",
        error_code="WEB_DOCUMENT_SOURCE_NOT_FOUND",
    )


def _image_source_not_found() -> dict[str, Any]:
    return envelope(
        False,
        "Không tìm thấy ảnh private đang hoạt động thuộc Web account hiện tại.",
        status_name="guarded",
        error_code="WEB_IMAGE_TO_PDF_SOURCE_NOT_FOUND",
    )


def _ocr_source_not_found() -> dict[str, Any]:
    return envelope(
        False,
        "Không tìm thấy ảnh private đang hoạt động thuộc Web account hiện tại.",
        status_name="guarded",
        error_code="WEB_DOCUMENT_OCR_SOURCE_NOT_FOUND",
    )


def _ocr_readiness_guard(exc: DocumentOperationError) -> dict[str, Any]:
    """Return an honest no-artifact state when optional local OCR is absent."""

    codes = {
        "OCR_RUNTIME_UNAVAILABLE": "WEB_DOCUMENT_OCR_RUNTIME_UNAVAILABLE",
        "OCR_LANGUAGE_UNAVAILABLE": "WEB_DOCUMENT_OCR_LANGUAGE_UNAVAILABLE",
    }
    return envelope(
        False,
        exc.public_message,
        status_name="guarded",
        error_code=codes.get(exc.code, "WEB_DOCUMENT_OCR_UNAVAILABLE"),
    )


def _operation_unavailable() -> dict[str, Any]:
    return envelope(
        False,
        "PDF đầu ra không còn sẵn sàng để tải. Hãy chạy thao tác mới hoặc liên hệ hỗ trợ.",
        status_name="guarded",
        error_code="WEB_DOCUMENT_OPERATION_UNAVAILABLE",
    )


def _record_event(conn, *, operation_id: str, state: str, when: str | None = None) -> None:
    if state not in OPERATION_STATES:
        raise RuntimeError("Trạng thái Document Operation không hợp lệ")
    row = conn.execute(
        "SELECT COALESCE(MAX(sequence), 0) + 1 FROM web_document_operation_events WHERE operation_id=?",
        (operation_id,),
    ).fetchone()
    sequence = int(row[0] or 1) if row else 1
    conn.execute(
        """INSERT INTO web_document_operation_events (id, operation_id, state, sequence, created_at)
           VALUES (?, ?, ?, ?, ?)""",
        (str(uuid.uuid4()), operation_id, state, sequence, when or utc_now()),
    )


def _request_fingerprint(*, source_asset_id: str, page_range: str, source_sha256: str, source_bytes: int) -> str:
    payload = json.dumps(
        {
            "kind": PDF_SPLIT_KIND,
            "source_asset_id": source_asset_id,
            "page_range": page_range,
            "source_sha256": source_sha256,
            "source_bytes": source_bytes,
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _merge_request_fingerprint(sources: list[dict[str, Any]]) -> str:
    """Bind a merge idempotency key to the ordered verified source set."""
    payload = json.dumps(
        {
            "kind": PDF_MERGE_KIND,
            "sources": [
                {
                    "asset_id": str(source["id"]),
                    "sha256": str(source["sha256"]),
                    "byte_size": int(source["byte_size"]),
                }
                for source in sources
            ],
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _image_to_pdf_request_fingerprint(sources: list[dict[str, Any]]) -> str:
    """Bind an idempotency key to the exact image order and source revisions."""
    payload = json.dumps(
        {
            "kind": IMAGE_TO_PDF_KIND,
            "sources": [
                {
                    "asset_id": str(source["id"]),
                    "sha256": str(source["sha256"]),
                    "byte_size": int(source["byte_size"]),
                }
                for source in sources
            ],
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _image_ocr_request_fingerprint(
    *,
    source_asset_id: str,
    source_sha256: str,
    source_bytes: int,
    language: str,
) -> str:
    """Bind server-side OCR replay to one private source revision and mode."""

    payload = json.dumps(
        {
            "kind": IMAGE_OCR_KIND,
            "source_asset_id": source_asset_id,
            "source_sha256": source_sha256,
            "source_bytes": source_bytes,
            "language": language,
            "engine": "local_tesseract_text_only",
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _optimize_request_fingerprint(*, source_asset_id: str, source_sha256: str, source_bytes: int) -> str:
    """Bind one structural optimize intent to its verified source revision."""
    payload = json.dumps(
        {
            "kind": PDF_OPTIMIZE_KIND,
            "source_asset_id": source_asset_id,
            "source_sha256": source_sha256,
            "source_bytes": source_bytes,
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _pdf_to_word_request_fingerprint(*, source_asset_id: str, source_sha256: str, source_bytes: int) -> str:
    """Bind PDF-text-to-DOCX replay to the exact verified source revision."""
    payload = json.dumps(
        {
            "kind": PDF_TO_WORD_KIND,
            "source_asset_id": source_asset_id,
            "source_sha256": source_sha256,
            "source_bytes": source_bytes,
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _pdf_to_images_request_fingerprint(*, source_asset_id: str, source_sha256: str, source_bytes: int) -> str:
    """Bind PDF-to-images replay to one exact verified PDF revision."""
    payload = json.dumps(
        {
            "kind": PDF_TO_IMAGES_KIND,
            "source_asset_id": source_asset_id,
            "source_sha256": source_sha256,
            "source_bytes": source_bytes,
            "render_scale": PDF_TO_IMAGES_RENDER_SCALE,
        },
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _output_spec(kind: str, *, output_page_count: int | None = None) -> tuple[str, str, str]:
    """Return server-owned suffix, MIME and attachment name for one kind."""
    if kind == PDF_TO_IMAGES_KIND and int(output_page_count or 0) == 1:
        return ".png", "image/png", "toan-aas-pdf-page-001.png"
    try:
        return OUTPUT_SPEC_BY_KIND[kind]
    except KeyError as exc:
        raise RuntimeError("Loại artifact Document Operation không hợp lệ") from exc


def _quota_available(conn, *, account_id: str, additional_bytes: int) -> bool:
    row = conn.execute(
        "SELECT COALESCE(SUM(byte_size), 0) FROM web_document_operations WHERE account_id=? AND byte_size IS NOT NULL",
        (account_id,),
    ).fetchone()
    used = int(row[0] or 0) if row else 0
    return used + additional_bytes <= _maximum_account_bytes()


def _operation_response(operation: dict[str, Any]) -> dict[str, Any]:
    state = str(operation.get("state") or "failed")
    kind = str(operation.get("kind") or PDF_SPLIT_KIND)
    failure_code = str(operation.get("_failure_code") or "")
    public_operation = {key: value for key, value in operation.items() if not key.startswith("_")}
    label, completed_message = {
        PDF_SPLIT_KIND: ("PDF Split", "Đã tách và xác minh PDF riêng tư."),
        PDF_MERGE_KIND: ("PDF Merge", "Đã gộp và xác minh PDF riêng tư."),
        PDF_OPTIMIZE_KIND: ("PDF Optimize", "Đã tối ưu và xác minh PDF riêng tư."),
        IMAGE_TO_PDF_KIND: ("Ảnh → PDF", "Đã tạo và xác minh PDF riêng tư từ ảnh."),
        PDF_TO_IMAGES_KIND: ("PDF → ảnh", "Đã render và xác minh PNG riêng tư từ PDF."),
        PDF_TO_WORD_KIND: ("PDF có text → Word", "Đã trích xuất text và xác minh DOCX riêng tư."),
        IMAGE_OCR_KIND: ("OCR ảnh", "Đã trích xuất và xác minh văn bản OCR private."),
    }.get(kind, ("Document Operation", "Đã xác minh artifact tài liệu riêng tư."))
    if state == "completed":
        return envelope(True, completed_message, data={"operation": public_operation}, status_name="completed")
    if state in {"queued", "processing"}:
        return envelope(True, f"{label} đang được máy chủ xử lý.", data={"operation": public_operation}, status_name=state)
    if state == "guarded" and kind == PDF_OPTIMIZE_KIND:
        return envelope(
            False,
            "Không có bản PDF nhỏ hơn đạt chuẩn an toàn; file gốc không thay đổi.",
            data={"operation": public_operation},
            status_name="guarded",
            error_code="WEB_DOCUMENT_OPERATION_NOT_REDUCED",
        )
    if state == "guarded" and kind == PDF_TO_WORD_KIND:
        return envelope(
            False,
            "PDF chưa có văn bản trích xuất được. File scan/ảnh cần OCR riêng; Web không phát DOCX giả.",
            data={"operation": public_operation},
            status_name="guarded",
            error_code="WEB_DOCUMENT_OPERATION_TEXT_NOT_FOUND",
        )
    if state == "guarded" and kind == IMAGE_OCR_KIND:
        if failure_code == "OCR_RUNTIME_UNAVAILABLE":
            return envelope(
                False,
                "OCR chưa có local runtime sẵn sàng để chạy lại an toàn. Không có file text được phát hành.",
                data={"operation": public_operation},
                status_name="guarded",
                error_code="WEB_DOCUMENT_OCR_RUNTIME_UNAVAILABLE",
            )
        if failure_code == "OCR_LANGUAGE_UNAVAILABLE":
            return envelope(
                False,
                "Language pack OCR đã chọn chưa sẵn sàng. Không có file text được phát hành.",
                data={"operation": public_operation},
                status_name="guarded",
                error_code="WEB_DOCUMENT_OCR_LANGUAGE_UNAVAILABLE",
            )
        return envelope(
            False,
            "OCR không đọc được văn bản đủ để phát hành. Web không tạo file text giả.",
            data={"operation": public_operation},
            status_name="guarded",
            error_code="WEB_DOCUMENT_OCR_TEXT_NOT_FOUND",
        )
    return envelope(False, f"{label} chưa thể hoàn tất an toàn.", data={"operation": public_operation}, status_name="failed", error_code="WEB_DOCUMENT_OPERATION_FAILED")


def _mark_failed(operation_id: str, account_id: str, *, kind: str, request: Request, code: str) -> None:
    if kind not in SUPPORTED_KINDS:
        raise RuntimeError("Loại Document Operation không hợp lệ")
    ensure_copyfast_schema()
    now = utc_now()
    with transaction() as conn:
        row = conn.execute(
            "SELECT state FROM web_document_operations WHERE id=? AND account_id=?",
            (operation_id, account_id),
        ).fetchone()
        if not row or str(row[0]) in {"completed", "unavailable", "guarded"}:
            return
        conn.execute(
            """UPDATE web_document_operations SET state='failed', failure_code=?, updated_at=?
               WHERE id=? AND account_id=?""",
            (code, now, operation_id, account_id),
        )
        _record_event(conn, operation_id=operation_id, state="failed", when=now)
        _record_audit(
            conn,
            account_id=account_id,
            canonical_user_id=None,
            action=f"web.document_operation.{kind}_failed",
            request_id=_request_id(request),
            target=operation_id,
            outcome="failed",
            detail=f"code={code}",
        )


def _mark_guarded(operation_id: str, account_id: str, *, kind: str, request: Request, code: str) -> None:
    """Record a safe terminal non-delivery without presenting it as failure/success."""
    if kind not in SUPPORTED_KINDS:
        raise RuntimeError("Loại Document Operation không hợp lệ")
    ensure_copyfast_schema()
    now = utc_now()
    with transaction() as conn:
        row = conn.execute(
            "SELECT state FROM web_document_operations WHERE id=? AND account_id=?",
            (operation_id, account_id),
        ).fetchone()
        if not row or str(row[0]) in {"completed", "unavailable", "guarded"}:
            return
        conn.execute(
            """UPDATE web_document_operations SET state='guarded', failure_code=?, updated_at=?
               WHERE id=? AND account_id=?""",
            (code, now, operation_id, account_id),
        )
        _record_event(conn, operation_id=operation_id, state="guarded", when=now)
        _record_audit(
            conn,
            account_id=account_id,
            canonical_user_id=None,
            action=f"web.document_operation.{kind}_guarded",
            request_id=_request_id(request),
            target=operation_id,
            outcome="guarded",
            detail=f"code={code}",
        )


def _mark_output_unavailable(operation_id: str, account_id: str) -> None:
    ensure_copyfast_schema()
    with transaction() as conn:
        updated = conn.execute(
            """UPDATE web_document_operations SET state='unavailable', updated_at=?
               WHERE id=? AND account_id=? AND state='completed'""",
            (utc_now(), operation_id, account_id),
        )
        if updated.rowcount:
            _record_event(conn, operation_id=operation_id, state="unavailable")


def _mark_sources_unavailable(asset_ids: list[str], account_id: str) -> None:
    unique_ids = sorted({str(asset_id) for asset_id in asset_ids if UUID_PATTERN.fullmatch(str(asset_id))})
    if not unique_ids:
        return
    with transaction() as conn:
        placeholders = ", ".join("?" for _ in unique_ids)
        conn.execute(
            f"""UPDATE web_asset_files
               SET state='unavailable', updated_at=?, lifecycle_revision=lifecycle_revision + 1
               WHERE id IN ({placeholders}) AND account_id=? AND state='active'""",
            (utc_now(), *unique_ids, account_id),
        )


def _mark_source_unavailable(asset_id: str, account_id: str) -> None:
    _mark_sources_unavailable([asset_id], account_id)


def _build_split_output(root: Path, source_copy: Path, *, page_range: str) -> tuple[Path, str, int, str, int, int, int]:
    """Create a stripped, verified PDF output from a bounded copied input."""
    temporary_output = _staging_path(root, ".pdf")
    final_path: Path | None = None
    try:
        PdfReader, PdfWriter = _pdf_classes()
        try:
            with source_copy.open("rb") as source_stream:
                reader = PdfReader(source_stream, strict=True)
                if reader.is_encrypted:
                    raise DocumentOperationError("PDF được mã hóa chưa thể tách an toàn", code="PDF_ENCRYPTED")
                source_page_count = len(reader.pages)
                indices, start_page, end_page = _selected_pages(page_range, source_page_count)
                writer = PdfWriter()
                # Start with a fresh writer, omit annotations/automatic actions
                # on every copied page, and write only neutral metadata. This
                # prevents interactive input PDF behavior from becoming part
                # of the delivered attachment.
                for index in indices:
                    writer.add_page(reader.pages[index], excluded_keys=PDF_EXCLUDED_PAGE_KEYS)
                writer.add_metadata({"/Title": "TOAN AAS PDF Split", "/Producer": "TOAN AAS Web PDF Split"})
                with temporary_output.open("xb") as output_stream:
                    writer.write(output_stream)
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF không hợp lệ hoặc không thể tách an toàn", code="PDF_PARSE_FAILED") from exc

        byte_size = temporary_output.stat().st_size
        if byte_size < 1 or byte_size > _maximum_output_bytes():
            raise DocumentOperationError("PDF đầu ra vượt giới hạn artifact an toàn", code="PDF_OUTPUT_LIMIT")
        digest = hashlib.sha256()
        with temporary_output.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
        output_digest = digest.hexdigest()
        try:
            with temporary_output.open("rb") as verify_stream:
                verified = PdfReader(verify_stream, strict=True)
                if verified.is_encrypted or len(verified.pages) != len(indices):
                    raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="PDF_OUTPUT_INVALID")
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="PDF_OUTPUT_INVALID") from exc

        outputs = _private_operation_directory(root, "outputs")
        storage_key = f"outputs/{uuid.uuid4().hex}.pdf"
        final_path = _output_path(root, storage_key)
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn PDF đầu ra không thuộc output storage riêng")
        os.replace(temporary_output, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=output_digest):
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra integrity", code="PDF_OUTPUT_INVALID")
        return final_path, storage_key, byte_size, output_digest, source_page_count, start_page, end_page
    except Exception:
        _safe_unlink(final_path)
        raise
    finally:
        _safe_unlink(temporary_output)


def _build_merge_output(root: Path, source_copies: list[Path]) -> tuple[Path, str, int, str, int]:
    """Merge bounded copied PDFs in explicit input order into a clean artifact."""
    if len(source_copies) < 2 or len(source_copies) > MAX_MERGE_SOURCES:
        raise DocumentOperationError("Cần từ 2 đến 8 PDF nguồn để gộp an toàn", code="PDF_MERGE_SOURCE_COUNT")
    temporary_output = _staging_path(root, ".pdf")
    final_path: Path | None = None
    try:
        PdfReader, PdfWriter = _pdf_classes()
        try:
            writer = PdfWriter()
            source_page_count = 0
            # Keep each private copied input open until writer.write(). This
            # avoids a lazy indirect-object read after an input file has been
            # closed while retaining a hard bound of eight files/30 pages.
            with ExitStack() as streams:
                for source_copy in source_copies:
                    source_stream = streams.enter_context(source_copy.open("rb"))
                    reader = PdfReader(source_stream, strict=True)
                    if reader.is_encrypted:
                        raise DocumentOperationError("PDF được mã hóa chưa thể gộp an toàn", code="PDF_ENCRYPTED")
                    page_count = len(reader.pages)
                    if page_count < 1:
                        raise DocumentOperationError("PDF nguồn không có trang hợp lệ để gộp", code="PDF_PAGE_LIMIT")
                    source_page_count += page_count
                    if source_page_count > MAX_PAGES:
                        raise DocumentOperationError(
                            f"Tổng số trang PDF vượt giới hạn {MAX_PAGES} trang/lần",
                            code="PDF_PAGE_LIMIT",
                        )
                    for page in reader.pages:
                        writer.add_page(page, excluded_keys=PDF_EXCLUDED_PAGE_KEYS)
                writer.add_metadata({"/Title": "TOAN AAS PDF Merge", "/Producer": "TOAN AAS Web PDF Merge"})
                with temporary_output.open("xb") as output_stream:
                    writer.write(output_stream)
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF không hợp lệ hoặc không thể gộp an toàn", code="PDF_PARSE_FAILED") from exc

        byte_size = temporary_output.stat().st_size
        if byte_size < 1 or byte_size > _maximum_output_bytes():
            raise DocumentOperationError("PDF đầu ra vượt giới hạn artifact an toàn", code="PDF_OUTPUT_LIMIT")
        digest = hashlib.sha256()
        with temporary_output.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
        output_digest = digest.hexdigest()
        try:
            with temporary_output.open("rb") as verify_stream:
                verified = PdfReader(verify_stream, strict=True)
                if verified.is_encrypted or len(verified.pages) != source_page_count:
                    raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="PDF_OUTPUT_INVALID")
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="PDF_OUTPUT_INVALID") from exc

        outputs = _private_operation_directory(root, "outputs")
        storage_key = f"outputs/{uuid.uuid4().hex}.pdf"
        final_path = _output_path(root, storage_key)
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn PDF đầu ra không thuộc output storage riêng")
        os.replace(temporary_output, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=output_digest):
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra integrity", code="PDF_OUTPUT_INVALID")
        return final_path, storage_key, byte_size, output_digest, source_page_count
    except Exception:
        _safe_unlink(final_path)
        raise
    finally:
        _safe_unlink(temporary_output)


def _image_format_matches(extension: str, image_format: str | None) -> bool:
    expected = {
        ".jpg": {"JPEG"},
        ".jpeg": {"JPEG"},
        ".png": {"PNG"},
        ".webp": {"WEBP"},
    }.get(extension, set())
    return str(image_format or "").upper() in expected


def _inspect_image_source(source_copy: Path, *, extension: str) -> int:
    """Verify the encoded image and return only its bounded raster size.

    This phase runs for every staged source before any page PDF is created,
    which means the aggregate pixel limit is enforced before a large decode or
    generated artifact is retained.
    """
    Image, ImageFile, _, UnidentifiedImageError = _image_classes()
    if ImageFile.LOAD_TRUNCATED_IMAGES:
        raise DocumentOperationError("Image runtime không ở chế độ kiểm tra đầy đủ", code="IMAGE_RUNTIME_UNAVAILABLE")
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(source_copy) as verifier:
                if not _image_format_matches(extension, verifier.format):
                    raise DocumentOperationError("Định dạng ảnh nguồn không khớp Asset Vault", code="IMAGE_SOURCE_INVALID")
                frame_count = int(getattr(verifier, "n_frames", 1) or 1)
                if frame_count != 1 or bool(getattr(verifier, "is_animated", False)):
                    raise DocumentOperationError("Ảnh động chưa được hỗ trợ để tạo PDF an toàn", code="IMAGE_ANIMATED")
                width, height = verifier.size
                pixels = int(width) * int(height)
                if width < 1 or height < 1:
                    raise DocumentOperationError("Kích thước ảnh không hợp lệ", code="IMAGE_DIMENSION_LIMIT")
                if width > MAX_IMAGE_DIMENSION or height > MAX_IMAGE_DIMENSION:
                    raise DocumentOperationError(
                        "Cạnh dài ảnh vượt giới hạn 7680 px mỗi nguồn",
                        code="IMAGE_DIMENSION_LIMIT",
                    )
                if max(width, height) / min(width, height) > MAX_IMAGE_ASPECT_RATIO:
                    raise DocumentOperationError(
                        "Tỷ lệ khung hình ảnh vượt giới hạn xử lý an toàn",
                        code="IMAGE_ASPECT_RATIO_LIMIT",
                    )
                if pixels > MAX_IMAGE_PIXELS_PER_SOURCE:
                    raise DocumentOperationError(
                        "Độ phân giải ảnh vượt giới hạn 16 MP mỗi nguồn",
                        code="IMAGE_PIXEL_LIMIT",
                    )
                # verify() checks the compressed source before a later full
                # decode. It does not retain raster data in memory.
                verifier.verify()
        return pixels
    except DocumentOperationError:
        raise
    except (Image.DecompressionBombWarning, Image.DecompressionBombError) as exc:
        raise DocumentOperationError(
            "Độ phân giải ảnh vượt giới hạn xử lý an toàn",
            code="IMAGE_PIXEL_LIMIT",
        ) from exc
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise DocumentOperationError("Ảnh không hợp lệ hoặc bị hỏng", code="IMAGE_PARSE_FAILED") from exc


def _render_image_page(source_copy: Path, page_pdf: Path, *, extension: str) -> int:
    """Decode, normalize and rasterize one trusted staging image into one PDF page.

    The input is already copied and hash-verified, but this step deliberately
    performs real decoder validation.  A fresh RGB image drops EXIF, comments,
    ICC profiles and transparency before the PDF writer sees any pixels.
    """
    Image, _, ImageOps, UnidentifiedImageError = _image_classes()
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            # Re-inspect immediately before decode. The first pass in the
            # batch builder enforces the aggregate bound; this second pass
            # preserves the same decoder checks at the point of use.
            pixels = _inspect_image_source(source_copy, extension=extension)

            with Image.open(source_copy) as decoded:
                if not _image_format_matches(extension, decoded.format):
                    raise DocumentOperationError("Định dạng ảnh nguồn không khớp Asset Vault", code="IMAGE_SOURCE_INVALID")
                if int(getattr(decoded, "n_frames", 1) or 1) != 1 or bool(getattr(decoded, "is_animated", False)):
                    raise DocumentOperationError("Ảnh động chưa được hỗ trợ để tạo PDF an toàn", code="IMAGE_ANIMATED")
                decoded.load()
                normalized = ImageOps.exif_transpose(decoded)
                rgba = None
                rgb = None
                alpha = None
                try:
                    rgba = normalized.convert("RGBA")
                    rgb = Image.new("RGB", rgba.size, (255, 255, 255))
                    alpha = rgba.getchannel("A")
                    rgb.paste(rgba, mask=alpha)
                    # Opening our own output handle with exclusive create
                    # avoids a destination race and keeps metadata opt-in.
                    with page_pdf.open("xb") as page_stream:
                        rgb.save(page_stream, format="PDF", resolution=144.0, quality=95, subsampling=0)
                finally:
                    if alpha is not None:
                        alpha.close()
                    if rgb is not None:
                        rgb.close()
                    if rgba is not None:
                        rgba.close()
                    if normalized is not decoded:
                        normalized.close()
        return pixels
    except DocumentOperationError:
        raise
    except (Image.DecompressionBombWarning, Image.DecompressionBombError) as exc:
        raise DocumentOperationError(
            "Độ phân giải ảnh vượt giới hạn xử lý an toàn",
            code="IMAGE_PIXEL_LIMIT",
        ) from exc
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise DocumentOperationError("Ảnh không hợp lệ hoặc bị hỏng", code="IMAGE_PARSE_FAILED") from exc
    except MemoryError as exc:
        raise DocumentOperationError("Ảnh vượt giới hạn bộ nhớ xử lý an toàn", code="IMAGE_RESOURCE_LIMIT") from exc


def _normalized_ocr_text(value: Any) -> str:
    """Keep a private OCR artifact readable, bounded and text-only.

    OCR may return arbitrary controls from malformed fonts.  Preserve normal
    newlines and tabs for a useful download, but remove other controls before
    the server creates an attachment.  The recognized text is never projected
    into an API response, audit event or browser state.
    """

    if not isinstance(value, str):
        raise DocumentOperationError("OCR không trả về văn bản hợp lệ", code="OCR_PARSE_FAILED")
    cleaned = value.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    cleaned = "".join(character for character in cleaned if character in {"\n", "\t"} or ord(character) >= 32)
    cleaned = "\n".join(line.rstrip() for line in cleaned.split("\n")).strip()
    if not cleaned:
        raise DocumentOperationError("OCR không tìm thấy văn bản để phát hành", code="OCR_TEXT_NOT_FOUND")
    if len(cleaned) > MAX_IMAGE_OCR_CHARACTERS:
        raise DocumentOperationError("Văn bản OCR vượt giới hạn artifact an toàn", code="OCR_OUTPUT_LIMIT")
    encoded = f"{cleaned}\n".encode("utf-8")
    if len(encoded) < 1 or len(encoded) > min(MAX_IMAGE_OCR_OUTPUT_BYTES, _maximum_output_bytes()):
        raise DocumentOperationError("Văn bản OCR vượt giới hạn artifact an toàn", code="OCR_OUTPUT_LIMIT")
    return cleaned + "\n"


def _build_image_ocr_output(
    root: Path,
    source_copy: Path,
    extension: str,
    *,
    language: str,
    ocr_module: Any,
) -> tuple[Path, str, int, str, int, int]:
    """Run local OCR over one inspected image and retain only verified text.

    This boundary deliberately has no remote URL, browser file, model choice,
    OCR result preview or generated-media behavior.  It receives a staged,
    integrity-checked Asset Vault copy and only publishes a private `.txt`
    after the text bytes have been re-read and hashed.
    """

    temporary_output = _staging_path(root, ".txt")
    final_path: Path | None = None
    try:
        Image, _, ImageOps, UnidentifiedImageError = _image_classes()
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            pixels = _inspect_image_source(source_copy, extension=extension)
            with Image.open(source_copy) as decoded:
                if not _image_format_matches(extension, decoded.format):
                    raise DocumentOperationError("Định dạng ảnh nguồn không khớp Asset Vault", code="IMAGE_SOURCE_INVALID")
                if int(getattr(decoded, "n_frames", 1) or 1) != 1 or bool(getattr(decoded, "is_animated", False)):
                    raise DocumentOperationError("Ảnh động chưa được hỗ trợ cho OCR an toàn", code="IMAGE_ANIMATED")
                decoded.load()
                normalized = ImageOps.exif_transpose(decoded)
                raster = None
                try:
                    raster = normalized.convert("RGB")
                    try:
                        raw_text = ocr_module.image_to_string(
                            raster,
                            lang=language,
                            config="--oem 1 --psm 6",
                            timeout=30,
                        )
                    except Exception as exc:
                        raise DocumentOperationError(
                            "Local Tesseract không thể hoàn tất OCR ảnh an toàn.",
                            code="OCR_RUNTIME_UNAVAILABLE",
                        ) from exc
                finally:
                    if raster is not None:
                        raster.close()
                    if normalized is not decoded:
                        normalized.close()

        text = _normalized_ocr_text(raw_text)
        encoded = text.encode("utf-8")
        with temporary_output.open("xb") as stream:
            stream.write(encoded)
        byte_size = temporary_output.stat().st_size
        if byte_size != len(encoded) or byte_size < 1 or byte_size > min(MAX_IMAGE_OCR_OUTPUT_BYTES, _maximum_output_bytes()):
            raise DocumentOperationError("Văn bản OCR vượt giới hạn artifact an toàn", code="OCR_OUTPUT_LIMIT")
        digest = hashlib.sha256(encoded).hexdigest()
        # Re-read/decode the exact staged artifact before publishing it. This
        # catches disk races/truncation and keeps the delivery contract UTF-8.
        if temporary_output.read_bytes() != encoded:
            raise DocumentOperationError("Văn bản OCR đầu ra không vượt qua kiểm tra", code="OCR_OUTPUT_INVALID")
        try:
            temporary_output.read_bytes().decode("utf-8", errors="strict")
        except UnicodeDecodeError as exc:
            raise DocumentOperationError("Văn bản OCR đầu ra không hợp lệ", code="OCR_OUTPUT_INVALID") from exc

        outputs = _private_operation_directory(root, "outputs")
        storage_key = f"outputs/{uuid.uuid4().hex}.txt"
        final_path = _output_path(root, storage_key, expected_suffix=".txt")
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn OCR đầu ra không thuộc output storage riêng")
        os.replace(temporary_output, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=digest):
            raise DocumentOperationError("Văn bản OCR đầu ra không vượt qua kiểm tra integrity", code="OCR_OUTPUT_INVALID")
        return final_path, storage_key, byte_size, digest, pixels, len(text.rstrip("\n"))
    except DocumentOperationError:
        _safe_unlink(final_path)
        raise
    except (Image.DecompressionBombWarning, Image.DecompressionBombError) as exc:
        raise DocumentOperationError("Độ phân giải ảnh vượt giới hạn xử lý an toàn", code="IMAGE_PIXEL_LIMIT") from exc
    except (UnidentifiedImageError, OSError, ValueError, MemoryError) as exc:
        raise DocumentOperationError("Không thể xử lý ảnh OCR an toàn", code="OCR_PARSE_FAILED") from exc
    finally:
        _safe_unlink(temporary_output)


def _build_image_to_pdf_output(
    root: Path,
    source_copies: list[tuple[Path, str]],
) -> tuple[Path, str, int, str, int, int]:
    """Build a clean one-page-per-image PDF from bounded private sources."""
    if not 1 <= len(source_copies) <= MAX_IMAGE_PDF_SOURCES:
        raise DocumentOperationError("Cần từ 1 đến 8 ảnh nguồn để tạo PDF an toàn", code="IMAGE_SOURCE_COUNT")
    temporary_output = _staging_path(root, ".pdf")
    page_pdfs: list[Path] = []
    final_path: Path | None = None
    try:
        total_pixels = 0
        # Inspect every encoded source before rendering any page. This makes
        # the aggregate resource decision early and avoids partially built
        # PDFs for a batch that is collectively too large.
        inspected_sources: list[tuple[Path, str]] = []
        for source_copy, extension in source_copies:
            pixels = _inspect_image_source(source_copy, extension=extension)
            total_pixels += pixels
            if total_pixels > MAX_IMAGE_PIXELS_TOTAL:
                raise DocumentOperationError(
                    "Tổng độ phân giải ảnh vượt giới hạn 32 MP mỗi lần",
                    code="IMAGE_TOTAL_PIXEL_LIMIT",
                )
            inspected_sources.append((source_copy, extension))
        for source_copy, extension in inspected_sources:
            page_pdf = _staging_path(root, ".image-page.pdf")
            page_pdfs.append(page_pdf)
            _render_image_page(source_copy, page_pdf, extension=extension)

        PdfReader, PdfWriter = _pdf_classes()
        try:
            writer = PdfWriter()
            with ExitStack() as streams:
                for page_pdf in page_pdfs:
                    page_stream = streams.enter_context(page_pdf.open("rb"))
                    generated = PdfReader(page_stream, strict=True)
                    if generated.is_encrypted or len(generated.pages) != 1:
                        raise DocumentOperationError("PDF trang ảnh không vượt qua kiểm tra", code="IMAGE_OUTPUT_INVALID")
                    writer.add_page(generated.pages[0], excluded_keys=PDF_EXCLUDED_PAGE_KEYS)
                # Keep every generated single-page PDF open until write():
                # pypdf may defer indirect-object reads from source streams.
                writer.add_metadata({"/Title": "TOAN AAS Image to PDF", "/Producer": "TOAN AAS Web Image to PDF"})
                with temporary_output.open("xb") as output_stream:
                    writer.write(output_stream)
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("Không thể tạo PDF từ ảnh an toàn", code="IMAGE_PARSE_FAILED") from exc

        byte_size = temporary_output.stat().st_size
        if byte_size < 1 or byte_size > _maximum_output_bytes():
            raise DocumentOperationError("PDF đầu ra vượt giới hạn artifact an toàn", code="PDF_OUTPUT_LIMIT")
        digest = hashlib.sha256()
        with temporary_output.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
        output_digest = digest.hexdigest()
        try:
            with temporary_output.open("rb") as verify_stream:
                verified = PdfReader(verify_stream, strict=True)
                if verified.is_encrypted or len(verified.pages) != len(source_copies):
                    raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="IMAGE_OUTPUT_INVALID")
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="IMAGE_OUTPUT_INVALID") from exc

        outputs = _private_operation_directory(root, "outputs")
        storage_key = f"outputs/{uuid.uuid4().hex}.pdf"
        final_path = _output_path(root, storage_key)
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn PDF đầu ra không thuộc output storage riêng")
        os.replace(temporary_output, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=output_digest):
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra integrity", code="IMAGE_OUTPUT_INVALID")
        return final_path, storage_key, byte_size, output_digest, len(source_copies), total_pixels
    except Exception:
        _safe_unlink(final_path)
        raise
    finally:
        _safe_unlink(temporary_output)
        for page_pdf in page_pdfs:
            _safe_unlink(page_pdf)


def _has_meaningful_optimization(*, source_bytes: int, output_bytes: int) -> bool:
    """Require a useful verified reduction before replacing a user's source."""
    if source_bytes < 1 or output_bytes < 1 or output_bytes >= source_bytes:
        return False
    saved_bytes = source_bytes - output_bytes
    # At least 1 KiB and 1% protects customers from a nominal byte-level
    # difference being labelled as a useful compression result.
    minimum_saved = max(
        MIN_OPTIMIZATION_SAVED_BYTES,
        int(source_bytes * MIN_OPTIMIZATION_SAVED_RATIO + 0.999999),
    )
    return saved_bytes >= minimum_saved


def _build_optimize_output(root: Path, source_copy: Path, *, source_bytes: int) -> tuple[Path, str, int, str, int]:
    """Create a structural PDF optimization only when it saves space."""
    temporary_output = _staging_path(root, ".pdf")
    final_path: Path | None = None
    try:
        PdfReader, PdfWriter = _pdf_classes()
        try:
            with source_copy.open("rb") as source_stream:
                reader = PdfReader(source_stream, strict=True)
                if reader.is_encrypted:
                    raise DocumentOperationError("PDF được mã hóa chưa thể tối ưu an toàn", code="PDF_ENCRYPTED")
                source_page_count = len(reader.pages)
                if source_page_count < 1 or source_page_count > MAX_PAGES:
                    raise DocumentOperationError(
                        f"PDF cần từ 1 đến {MAX_PAGES} trang để tối ưu an toàn",
                        code="PDF_PAGE_LIMIT",
                    )
                writer = PdfWriter()
                # A fresh writer deliberately retains the visual page content
                # while discarding page-level interactive behavior and source
                # metadata. It never downscales images or claims a lossy tier.
                for page in reader.pages:
                    writer.add_page(page, excluded_keys=PDF_EXCLUDED_PAGE_KEYS)
                for output_page in writer.pages:
                    output_page.compress_content_streams(level=9)
                writer.compress_identical_objects(remove_duplicates=True, remove_unreferenced=True)
                writer.add_metadata({"/Title": "TOAN AAS PDF Optimize", "/Producer": "TOAN AAS Web PDF Optimize"})
                with temporary_output.open("xb") as output_stream:
                    writer.write(output_stream)
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF không hợp lệ hoặc không thể tối ưu cấu trúc an toàn", code="PDF_PARSE_FAILED") from exc

        byte_size = temporary_output.stat().st_size
        if byte_size < 1 or byte_size > _maximum_output_bytes():
            raise DocumentOperationError("PDF đầu ra vượt giới hạn artifact an toàn", code="PDF_OUTPUT_LIMIT")
        if not _has_meaningful_optimization(source_bytes=source_bytes, output_bytes=byte_size):
            raise DocumentOperationError(
                "Không có bản PDF nhỏ hơn đạt chuẩn an toàn; file gốc không thay đổi",
                code="PDF_NOT_REDUCED",
            )
        digest = hashlib.sha256()
        with temporary_output.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
        output_digest = digest.hexdigest()
        try:
            with temporary_output.open("rb") as verify_stream:
                verified = PdfReader(verify_stream, strict=True)
                if verified.is_encrypted or len(verified.pages) != source_page_count:
                    raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="PDF_OUTPUT_INVALID")
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra", code="PDF_OUTPUT_INVALID") from exc

        outputs = _private_operation_directory(root, "outputs")
        storage_key = f"outputs/{uuid.uuid4().hex}.pdf"
        final_path = _output_path(root, storage_key)
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn PDF đầu ra không thuộc output storage riêng")
        os.replace(temporary_output, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=output_digest):
            raise DocumentOperationError("PDF đầu ra không vượt qua kiểm tra integrity", code="PDF_OUTPUT_INVALID")
        return final_path, storage_key, byte_size, output_digest, source_page_count
    except Exception:
        _safe_unlink(final_path)
        raise
    finally:
        _safe_unlink(temporary_output)


def _safe_docx_text(value: str) -> str:
    """Keep only XML 1.0 characters python-docx can safely serialize."""
    normalized = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return "".join(
        character
        for character in normalized
        if character in {"\t", "\n"}
        or 0x20 <= ord(character) <= 0xD7FF
        or 0xE000 <= ord(character) <= 0xFFFD
        or 0x10000 <= ord(character) <= 0x10FFFF
    )


def _text_paragraphs(value: str) -> list[str]:
    """Turn extracted logical text into bounded DOCX paragraphs without OCR."""
    return [line.strip() for line in _safe_docx_text(value).split("\n") if line.strip()]


def _verify_docx_output(path: Path, *, expected_paragraphs: list[str], expected_bytes: int) -> None:
    """Reject an unsafe/incomplete generated DOCX before it reaches outputs.

    The DOCX must be the document that this process just generated: bounded,
    ordinary OOXML only, no macro/ActiveX/embed payload, no ZIP traversal or
    external relationship, and its reopened visible paragraphs must match the
    bounded text extraction used to create it.
    """
    if expected_bytes < 1 or expected_bytes > _maximum_output_bytes():
        raise DocumentOperationError("DOCX đầu ra vượt giới hạn artifact an toàn", code="PDF_TO_WORD_OUTPUT_LIMIT")
    if not expected_paragraphs or len(expected_paragraphs) > MAX_PDF_TO_WORD_PARAGRAPHS:
        raise DocumentOperationError("DOCX đầu ra không vượt qua kiểm tra text", code="PDF_TO_WORD_OUTPUT_INVALID")
    try:
        with ZipFile(path) as archive:
            infos = archive.infolist()
            if not infos or len(infos) > MAX_DOCX_ARCHIVE_MEMBERS:
                raise DocumentOperationError("DOCX đầu ra có cấu trúc không an toàn", code="PDF_TO_WORD_OUTPUT_INVALID")
            names: set[str] = set()
            uncompressed_bytes = 0
            for info in infos:
                name = str(info.filename or "")
                posix_name = name.replace("\\", "/")
                parts = [part for part in posix_name.split("/") if part]
                mode = (info.external_attr >> 16) & 0o170000
                if (
                    not name
                    or name != posix_name
                    or posix_name.startswith("/")
                    or ".." in parts
                    or info.flag_bits & 0x1
                    or mode == 0o120000
                ):
                    raise DocumentOperationError("DOCX đầu ra có cấu trúc không an toàn", code="PDF_TO_WORD_OUTPUT_INVALID")
                lowered = posix_name.lower()
                if (
                    "vbaproject" in lowered
                    or lowered.startswith("word/embeddings/")
                    or lowered.startswith("word/activex/")
                ):
                    raise DocumentOperationError("DOCX đầu ra chứa thành phần không được phép", code="PDF_TO_WORD_OUTPUT_INVALID")
                uncompressed_bytes += max(0, int(info.file_size))
                if uncompressed_bytes > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise DocumentOperationError("DOCX đầu ra vượt giới hạn an toàn", code="PDF_TO_WORD_OUTPUT_LIMIT")
                names.add(posix_name)
                if lowered.endswith(".rels") and b'TargetMode="External"' in archive.read(info):
                    raise DocumentOperationError("DOCX đầu ra có liên kết ngoài không được phép", code="PDF_TO_WORD_OUTPUT_INVALID")
            if {"[Content_Types].xml", "word/document.xml"} - names:
                raise DocumentOperationError("DOCX đầu ra thiếu cấu trúc bắt buộc", code="PDF_TO_WORD_OUTPUT_INVALID")
    except DocumentOperationError:
        raise
    except (BadZipFile, OSError) as exc:
        raise DocumentOperationError("DOCX đầu ra không vượt qua kiểm tra", code="PDF_TO_WORD_OUTPUT_INVALID") from exc

    try:
        Document = _word_classes()
        reopened = Document(str(path))
        actual_paragraphs = [text for paragraph in reopened.paragraphs if (text := _safe_docx_text(paragraph.text)).strip()]
    except Exception as exc:
        raise DocumentOperationError("DOCX đầu ra không thể mở lại an toàn", code="PDF_TO_WORD_OUTPUT_INVALID") from exc
    if actual_paragraphs != expected_paragraphs:
        raise DocumentOperationError("DOCX đầu ra không vượt qua kiểm tra text", code="PDF_TO_WORD_OUTPUT_INVALID")


def _build_pdf_to_word_output(root: Path, source_copy: Path) -> tuple[Path, str, int, str, int, int, int]:
    """Export real extracted PDF text into a freshly generated private DOCX.

    This deliberately reads no source attachments, links, annotations, images,
    metadata or layout. It creates a new document from text only, so a scan or
    PDF without selectable text remains guarded instead of receiving a fake
    blank DOCX or an unstated OCR/provider fallback.
    """
    temporary_output = _staging_path(root, ".docx")
    final_path: Path | None = None
    try:
        PdfReader, _ = _pdf_classes()
        Document = _word_classes()
        extracted_paragraphs: list[str] = []
        source_page_count = 0
        extracted_characters = 0
        try:
            with source_copy.open("rb") as source_stream:
                reader = PdfReader(source_stream, strict=True)
                if reader.is_encrypted:
                    raise DocumentOperationError("PDF được mã hóa chưa thể trích xuất text an toàn", code="PDF_ENCRYPTED")
                source_page_count = len(reader.pages)
                if source_page_count < 1 or source_page_count > MAX_PAGES:
                    raise DocumentOperationError(
                        f"PDF cần từ 1 đến {MAX_PAGES} trang để trích xuất text an toàn",
                        code="PDF_PAGE_LIMIT",
                    )
                for page in reader.pages:
                    try:
                        raw_text = page.extract_text() or ""
                    except Exception as exc:
                        raise DocumentOperationError(
                            "Không thể trích xuất văn bản từ PDF này an toàn",
                            code="PDF_TEXT_EXTRACTION_FAILED",
                        ) from exc
                    safe_text = _safe_docx_text(raw_text)
                    page_characters = len(safe_text)
                    if page_characters > MAX_PDF_TO_WORD_PAGE_CHARACTERS:
                        raise DocumentOperationError(
                            "Văn bản một trang PDF vượt giới hạn xử lý an toàn",
                            code="PDF_TEXT_LIMIT",
                        )
                    extracted_characters += page_characters
                    if extracted_characters > MAX_PDF_TO_WORD_CHARACTERS:
                        raise DocumentOperationError(
                            "Văn bản PDF vượt giới hạn xử lý an toàn",
                            code="PDF_TEXT_LIMIT",
                        )
                    extracted_paragraphs.extend(_text_paragraphs(safe_text))
                    if len(extracted_paragraphs) > MAX_PDF_TO_WORD_PARAGRAPHS:
                        raise DocumentOperationError(
                            "PDF tạo quá nhiều đoạn văn để xuất an toàn",
                            code="PDF_TEXT_LIMIT",
                        )
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF không hợp lệ hoặc không thể trích xuất text an toàn", code="PDF_PARSE_FAILED") from exc

        if not extracted_paragraphs:
            raise DocumentOperationError(
                "PDF chưa có văn bản trích xuất được. File scan/ảnh cần OCR riêng; Web không phát DOCX giả.",
                code="PDF_TEXT_NOT_FOUND",
            )

        document = Document()
        document.core_properties.title = "TOAN AAS PDF text export"
        document.core_properties.author = "TOAN AAS Web"
        document.core_properties.subject = "Extracted PDF text"
        for index, paragraph_text in enumerate(extracted_paragraphs):
            document.add_paragraph(paragraph_text)
        document.save(str(temporary_output))

        byte_size = temporary_output.stat().st_size
        _verify_docx_output(
            temporary_output,
            expected_paragraphs=extracted_paragraphs,
            expected_bytes=byte_size,
        )
        digest = hashlib.sha256()
        with temporary_output.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
        output_digest = digest.hexdigest()
        outputs = _private_operation_directory(root, "outputs")
        storage_key = f"outputs/{uuid.uuid4().hex}.docx"
        final_path = _output_path(root, storage_key, expected_suffix=".docx")
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn DOCX đầu ra không thuộc output storage riêng")
        os.replace(temporary_output, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=output_digest):
            raise DocumentOperationError("DOCX đầu ra không vượt qua kiểm tra integrity", code="PDF_TO_WORD_OUTPUT_INVALID")
        return (
            final_path,
            storage_key,
            byte_size,
            output_digest,
            source_page_count,
            extracted_characters,
            len(extracted_paragraphs),
        )
    except Exception:
        _safe_unlink(final_path)
        raise
    finally:
        _safe_unlink(temporary_output)


def _verify_pdf_to_images_png_bytes(
    payload: bytes,
    *,
    expected_width: int | None = None,
    expected_height: int | None = None,
    expected_digest: str | None = None,
) -> tuple[int, str, int, int]:
    """Strictly parse one renderer-produced PNG without trusting its filename.

    The renderer creates these bytes, but delivery still requires a fresh
    decoder check. This detects a truncated/stale staging write before it can
    become either a private PNG attachment or a ZIP member.
    """
    Image, ImageFile, _, UnidentifiedImageError = _image_classes()
    byte_size = len(payload)
    if byte_size < 1 or byte_size > MAX_PDF_TO_IMAGES_PAGE_PNG_BYTES:
        raise DocumentOperationError("PNG một trang vượt giới hạn artifact an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
    if not payload.startswith(b"\x89PNG\r\n\x1a\n"):
        raise DocumentOperationError("PNG đầu ra không vượt qua kiểm tra", code="PDF_TO_IMAGES_OUTPUT_INVALID")
    if ImageFile.LOAD_TRUNCATED_IMAGES:
        raise DocumentOperationError("Image runtime không ở chế độ kiểm tra đầy đủ", code="PDF_TO_IMAGES_RUNTIME_UNAVAILABLE")
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(BytesIO(payload)) as verifier:
                if str(verifier.format or "").upper() != "PNG":
                    raise DocumentOperationError("PNG đầu ra không đúng định dạng", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                verifier.verify()
            with Image.open(BytesIO(payload)) as decoded:
                if str(decoded.format or "").upper() != "PNG":
                    raise DocumentOperationError("PNG đầu ra không đúng định dạng", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                width, height = (int(decoded.size[0]), int(decoded.size[1]))
                pixels = width * height
                if width < 1 or height < 1 or width > MAX_PDF_TO_IMAGES_DIMENSION or height > MAX_PDF_TO_IMAGES_DIMENSION:
                    raise DocumentOperationError("Kích thước PNG đầu ra vượt giới hạn an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
                if pixels > MAX_PDF_TO_IMAGES_PAGE_PIXELS:
                    raise DocumentOperationError("Độ phân giải PNG đầu ra vượt giới hạn an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
                if str(decoded.mode or "") != "RGB":
                    raise DocumentOperationError("PNG đầu ra có pixel format không hợp lệ", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                if expected_width is not None and width != expected_width:
                    raise DocumentOperationError("PNG đầu ra không khớp kích thước render", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                if expected_height is not None and height != expected_height:
                    raise DocumentOperationError("PNG đầu ra không khớp kích thước render", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                decoded.load()
    except DocumentOperationError:
        raise
    except (Image.DecompressionBombWarning, Image.DecompressionBombError) as exc:
        raise DocumentOperationError("Độ phân giải PNG đầu ra vượt giới hạn an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT") from exc
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise DocumentOperationError("PNG đầu ra không vượt qua kiểm tra", code="PDF_TO_IMAGES_OUTPUT_INVALID") from exc
    digest = hashlib.sha256(payload).hexdigest()
    if expected_digest is not None and not hmac.compare_digest(digest, expected_digest):
        raise DocumentOperationError("PNG đầu ra không vượt qua kiểm tra integrity", code="PDF_TO_IMAGES_OUTPUT_INVALID")
    return byte_size, digest, width, height


def _pdf_to_images_geometry(page) -> tuple[int, int, int]:
    """Bound a PDFium page before allocating its 2× bitmap."""
    try:
        source_width, source_height = page.get_size()
        width = int(math.ceil(float(source_width) * PDF_TO_IMAGES_RENDER_SCALE))
        height = int(math.ceil(float(source_height) * PDF_TO_IMAGES_RENDER_SCALE))
    except (AttributeError, TypeError, ValueError, OverflowError) as exc:
        raise DocumentOperationError("Kích thước trang PDF không hợp lệ", code="PDF_RENDER_DIMENSION_LIMIT") from exc
    if width < 1 or height < 1 or width > MAX_PDF_TO_IMAGES_DIMENSION or height > MAX_PDF_TO_IMAGES_DIMENSION:
        raise DocumentOperationError(
            f"Kích thước render PDF vượt giới hạn {MAX_PDF_TO_IMAGES_DIMENSION}px mỗi cạnh",
            code="PDF_RENDER_DIMENSION_LIMIT",
        )
    pixels = width * height
    if pixels > MAX_PDF_TO_IMAGES_PAGE_PIXELS:
        raise DocumentOperationError(
            "Độ phân giải render PDF vượt giới hạn 8 MP mỗi trang",
            code="PDF_RENDER_PIXEL_LIMIT",
        )
    return width, height, pixels


def _render_pdf_to_images_page(pdf_document, *, page_index: int, destination: Path, expected_width: int, expected_height: int) -> dict[str, Any]:
    """Render one preflighted PDF page into a metadata-free verified PNG."""
    page = None
    bitmap = None
    rendered = None
    normalized = None
    try:
        page = pdf_document[page_index]
        actual_width, actual_height, _ = _pdf_to_images_geometry(page)
        if actual_width != expected_width or actual_height != expected_height:
            raise DocumentOperationError("Kích thước trang PDF thay đổi trong khi render", code="PDF_RENDER_DIMENSION_LIMIT")
        bitmap = page.render(scale=PDF_TO_IMAGES_RENDER_SCALE)
        rendered = bitmap.to_pil()
        normalized = rendered.convert("RGB")
        if tuple(normalized.size) != (expected_width, expected_height):
            raise DocumentOperationError("PNG render không khớp kích thước trang PDF", code="PDF_TO_IMAGES_OUTPUT_INVALID")
        # `xb` prevents a stale or attacker-created staging filename from
        # being overwritten. Conversion creates fresh RGB pixels and does not
        # carry source PDF metadata, links, attachments or annotations.
        with destination.open("xb") as output_stream:
            normalized.save(output_stream, format="PNG", optimize=True)
        byte_size, digest, width, height = _verify_pdf_to_images_png_bytes(
            destination.read_bytes(),
            expected_width=expected_width,
            expected_height=expected_height,
        )
        return {
            "name": f"page_{page_index + 1:03d}.png",
            "path": destination,
            "byte_size": byte_size,
            "sha256": digest,
            "width": width,
            "height": height,
        }
    except DocumentOperationError:
        raise
    except MemoryError as exc:
        raise DocumentOperationError("PDF vượt giới hạn bộ nhớ render an toàn", code="PDF_RENDER_RESOURCE_LIMIT") from exc
    except Exception as exc:
        raise DocumentOperationError("Không thể render trang PDF an toàn", code="PDF_RENDER_FAILED") from exc
    finally:
        for resource in (normalized, rendered, bitmap, page):
            if resource is not None:
                try:
                    resource.close()
                except Exception:
                    pass


def _verify_pdf_to_images_zip(path: Path, expected_pages: list[dict[str, Any]]) -> tuple[int, str]:
    """Re-open the final ZIP and verify each expected private PNG member."""
    try:
        byte_size = path.stat().st_size
    except OSError as exc:
        raise DocumentOperationError("ZIP PNG đầu ra không còn sẵn sàng", code="PDF_TO_IMAGES_OUTPUT_INVALID") from exc
    if byte_size < 1 or byte_size > _maximum_output_bytes():
        raise DocumentOperationError("ZIP PNG đầu ra vượt giới hạn artifact an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
    if not 2 <= len(expected_pages) <= MAX_PDF_TO_IMAGES_ZIP_MEMBERS:
        raise DocumentOperationError("ZIP PNG có số trang không hợp lệ", code="PDF_TO_IMAGES_OUTPUT_INVALID")
    total_uncompressed = 0
    try:
        with ZipFile(path, "r") as archive:
            if archive.comment:
                raise DocumentOperationError("ZIP PNG đầu ra có metadata không được phép", code="PDF_TO_IMAGES_OUTPUT_INVALID")
            infos = archive.infolist()
            if len(infos) != len(expected_pages):
                raise DocumentOperationError("ZIP PNG đầu ra thiếu hoặc thừa trang", code="PDF_TO_IMAGES_OUTPUT_INVALID")
            for info, expected in zip(infos, expected_pages):
                expected_name = str(expected["name"])
                if info.is_dir() or info.filename != expected_name or info.compress_type != ZIP_DEFLATED:
                    raise DocumentOperationError("ZIP PNG đầu ra có entry không hợp lệ", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                if info.date_time != PDF_TO_IMAGES_ZIP_DATE_TIME or info.extra or info.comment:
                    raise DocumentOperationError("ZIP PNG đầu ra có metadata không được phép", code="PDF_TO_IMAGES_OUTPUT_INVALID")
                if info.file_size != int(expected["byte_size"]) or info.file_size > MAX_PDF_TO_IMAGES_PAGE_PNG_BYTES:
                    raise DocumentOperationError("ZIP PNG đầu ra vượt giới hạn một trang", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
                total_uncompressed += int(info.file_size)
                if total_uncompressed > MAX_PDF_TO_IMAGES_TOTAL_PNG_BYTES:
                    raise DocumentOperationError("ZIP PNG đầu ra vượt giới hạn ảnh đã render", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
                payload = archive.read(info)
                _verify_pdf_to_images_png_bytes(
                    payload,
                    expected_width=int(expected["width"]),
                    expected_height=int(expected["height"]),
                    expected_digest=str(expected["sha256"]),
                )
    except DocumentOperationError:
        raise
    except (BadZipFile, OSError, RuntimeError) as exc:
        raise DocumentOperationError("ZIP PNG đầu ra không vượt qua kiểm tra", code="PDF_TO_IMAGES_OUTPUT_INVALID") from exc
    digest = hashlib.sha256()
    try:
        with path.open("rb") as stream:
            while True:
                chunk = stream.read(CHUNK_BYTES)
                if not chunk:
                    break
                digest.update(chunk)
    except OSError as exc:
        raise DocumentOperationError("ZIP PNG đầu ra không còn sẵn sàng", code="PDF_TO_IMAGES_OUTPUT_INVALID") from exc
    return byte_size, digest.hexdigest()


def _build_pdf_to_images_output(root: Path, source_copy: Path) -> tuple[Path, str, int, str, int, int, str, str]:
    """Render all bounded PDF pages at Bot-compatible 2× into private delivery.

    A single page is delivered as a verified PNG. Multiple pages keep Bot's
    ZIP behavior but are stored privately and can only be fetched through the
    signed-session attachment route.
    """
    page_paths: list[Path] = []
    final_path: Path | None = None
    temporary_zip: Path | None = None
    pdf_document = None
    try:
        PdfReader, _ = _pdf_classes()
        try:
            with source_copy.open("rb") as source_stream:
                reader = PdfReader(source_stream, strict=True)
                if reader.is_encrypted:
                    raise DocumentOperationError("PDF được mã hóa chưa thể render ảnh an toàn", code="PDF_ENCRYPTED")
                source_page_count = len(reader.pages)
        except DocumentOperationError:
            raise
        except Exception as exc:
            raise DocumentOperationError("PDF không hợp lệ hoặc không thể đọc an toàn", code="PDF_PARSE_FAILED") from exc
        if source_page_count < 1 or source_page_count > MAX_PAGES:
            raise DocumentOperationError(
                f"PDF cần từ 1 đến {MAX_PAGES} trang để render ảnh an toàn",
                code="PDF_PAGE_LIMIT",
            )

        pdfium = _pdf_to_images_classes()
        try:
            pdf_document = pdfium.PdfDocument(str(source_copy), autoclose=False)
            if len(pdf_document) != source_page_count:
                raise DocumentOperationError("Số trang PDF không nhất quán", code="PDF_PARSE_FAILED")
            geometry: list[tuple[int, int]] = []
            total_pixels = 0
            for page_index in range(source_page_count):
                page = pdf_document[page_index]
                try:
                    width, height, pixels = _pdf_to_images_geometry(page)
                finally:
                    try:
                        page.close()
                    except Exception:
                        pass
                total_pixels += pixels
                if total_pixels > MAX_PDF_TO_IMAGES_TOTAL_PIXELS:
                    raise DocumentOperationError(
                        "Tổng độ phân giải render PDF vượt giới hạn 48 MP mỗi lần",
                        code="PDF_RENDER_TOTAL_PIXEL_LIMIT",
                    )
                geometry.append((width, height))
        except DocumentOperationError:
            raise
        except MemoryError as exc:
            raise DocumentOperationError("PDF vượt giới hạn bộ nhớ render an toàn", code="PDF_RENDER_RESOURCE_LIMIT") from exc
        except Exception as exc:
            raise DocumentOperationError("PDF không thể mở bằng renderer an toàn", code="PDF_PARSE_FAILED") from exc

        rendered_pages: list[dict[str, Any]] = []
        total_png_bytes = 0
        for page_index, (width, height) in enumerate(geometry):
            page_path = _staging_path(root, ".pdf-page.png")
            page_paths.append(page_path)
            rendered = _render_pdf_to_images_page(
                pdf_document,
                page_index=page_index,
                destination=page_path,
                expected_width=width,
                expected_height=height,
            )
            total_png_bytes += int(rendered["byte_size"])
            if total_png_bytes > MAX_PDF_TO_IMAGES_TOTAL_PNG_BYTES:
                raise DocumentOperationError("Tổng PNG render vượt giới hạn artifact an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
            rendered_pages.append(rendered)

        outputs = _private_operation_directory(root, "outputs")
        if source_page_count == 1:
            page = rendered_pages[0]
            byte_size = int(page["byte_size"])
            if byte_size > _maximum_output_bytes():
                raise DocumentOperationError("PNG đầu ra vượt giới hạn artifact an toàn", code="PDF_TO_IMAGES_OUTPUT_LIMIT")
            storage_key = f"outputs/{uuid.uuid4().hex}.png"
            final_path = _output_path(root, storage_key, expected_suffix=".png")
            if final_path.parent != outputs:
                raise RuntimeError("Đường dẫn PNG đầu ra không thuộc output storage riêng")
            os.replace(Path(page["path"]), final_path)
            if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=str(page["sha256"])):
                raise DocumentOperationError("PNG đầu ra không vượt qua kiểm tra integrity", code="PDF_TO_IMAGES_OUTPUT_INVALID")
            return (
                final_path,
                storage_key,
                byte_size,
                str(page["sha256"]),
                source_page_count,
                source_page_count,
                "image/png",
                "toan-aas-pdf-page-001.png",
            )

        temporary_zip = _staging_path(root, ".pdf-pages.zip")
        try:
            with ZipFile(temporary_zip, "x", compression=ZIP_DEFLATED, compresslevel=9) as archive:
                for page in rendered_pages:
                    # `write()` carries the staging file's mtime into the
                    # archive.  A fixed ZipInfo keeps this private multi-page
                    # result deterministic without exposing host timestamps.
                    member = ZipInfo(filename=str(page["name"]), date_time=PDF_TO_IMAGES_ZIP_DATE_TIME)
                    member.create_system = 3
                    member.external_attr = 0o600 << 16
                    member.compress_type = ZIP_DEFLATED
                    archive.writestr(
                        member,
                        Path(page["path"]).read_bytes(),
                        compress_type=ZIP_DEFLATED,
                        compresslevel=9,
                    )
        except OSError as exc:
            raise DocumentOperationError("Không thể đóng gói PNG PDF an toàn", code="PDF_TO_IMAGES_OUTPUT_INVALID") from exc
        byte_size, output_digest = _verify_pdf_to_images_zip(temporary_zip, rendered_pages)
        storage_key = f"outputs/{uuid.uuid4().hex}.zip"
        final_path = _output_path(root, storage_key, expected_suffix=".zip")
        if final_path.parent != outputs:
            raise RuntimeError("Đường dẫn ZIP đầu ra không thuộc output storage riêng")
        os.replace(temporary_zip, final_path)
        if not _verify_file(final_path, expected_bytes=byte_size, expected_digest=output_digest):
            raise DocumentOperationError("ZIP PNG đầu ra không vượt qua kiểm tra integrity", code="PDF_TO_IMAGES_OUTPUT_INVALID")
        return (
            final_path,
            storage_key,
            byte_size,
            output_digest,
            source_page_count,
            source_page_count,
            "application/zip",
            "toan-aas-pdf-pages.zip",
        )
    except Exception:
        _safe_unlink(final_path)
        raise
    finally:
        if pdf_document is not None:
            try:
                pdf_document.close()
            except Exception:
                pass
        _safe_unlink(temporary_zip)
        for page_path in page_paths:
            _safe_unlink(page_path)


def reconcile_document_operation_storage() -> None:
    """Remove only old unreferenced staging/generated outputs after interruption."""
    if not document_operations_enabled():
        return
    ensure_copyfast_schema()
    root = document_operations_directory()
    staging = _private_operation_directory(root, ".staging")
    outputs = _private_operation_directory(root, "outputs")
    cutoff_at = (datetime.now(timezone.utc) - timedelta(seconds=ORPHAN_RETENTION_SECONDS)).isoformat(timespec="seconds")
    now = utc_now()
    with transaction() as conn:
        referenced = {str(row[0]) for row in conn.execute("SELECT storage_key FROM web_document_operations WHERE storage_key IS NOT NULL").fetchall()}
        # An interrupted synchronous parser must never remain replayable as
        # `queued`/`processing` forever. Wait one retention window so a
        # short application restart cannot race an active worker, then record
        # an explicit failed lifecycle without inventing an output.
        interrupted = conn.execute(
            """SELECT id FROM web_document_operations
               WHERE state IN ('queued', 'processing') AND updated_at < ?""",
            (cutoff_at,),
        ).fetchall()
        for row in interrupted:
            operation_id = str(row[0])
            conn.execute(
                """UPDATE web_document_operations
                   SET state='failed', failure_code='INTERRUPTED', updated_at=?
                   WHERE id=? AND state IN ('queued', 'processing')""",
                (now, operation_id),
            )
            _record_event(conn, operation_id=operation_id, state="failed", when=now)
    cutoff = datetime.now(timezone.utc).timestamp() - ORPHAN_RETENTION_SECONDS
    for directory, check_reference in ((staging, False), (outputs, True)):
        try:
            candidates = list(directory.iterdir())
        except OSError:
            continue
        for candidate in candidates:
            try:
                if not candidate.is_file() or candidate.is_symlink() or candidate.stat().st_mtime > cutoff:
                    continue
                relative = candidate.resolve().relative_to(root.resolve()).as_posix()
            except (OSError, ValueError):
                continue
            if check_reference and relative in referenced:
                continue
            _safe_unlink(candidate)


@router.get("")
async def list_document_operations(
    limit: int = 50,
    kind: str | None = None,
    offset: int = Query(0, ge=0, le=10000),
    account: dict = Depends(require_account),
):
    _require_enabled()
    bounded_limit = max(1, min(int(limit), 100))
    requested_kind = str(kind or "").strip().lower()
    if requested_kind and requested_kind not in SUPPORTED_KINDS:
        raise HTTPException(status_code=422, detail="Loại Document Operation không hợp lệ")
    ensure_copyfast_schema()
    with transaction() as conn:
        if requested_kind:
            rows = conn.execute(
                f"""SELECT {OPERATION_SELECT} FROM web_document_operations
                    WHERE account_id=? AND kind=? ORDER BY updated_at DESC, id DESC LIMIT ? OFFSET ?""",
                (str(account["id"]), requested_kind, bounded_limit + 1, int(offset)),
            ).fetchall()
        else:
            rows = conn.execute(
                f"""SELECT {OPERATION_SELECT} FROM web_document_operations
                    WHERE account_id=? ORDER BY updated_at DESC, id DESC LIMIT ? OFFSET ?""",
                (str(account["id"]), bounded_limit + 1, int(offset)),
            ).fetchall()
    has_more = len(rows) > bounded_limit
    items: list[dict[str, Any]] = []
    for row in rows[:bounded_limit]:
        # `_failure_code` is a private branch helper for the one-operation
        # response. A history list must not disclose internal failure detail.
        public_operation = _operation_public(tuple(row))
        public_operation.pop("_failure_code", None)
        items.append(public_operation)
    return envelope(
        True,
        "Đã tải thao tác tài liệu Web.",
        data={
            "items": items,
            "has_more": has_more,
            "next_offset": int(offset) + bounded_limit if has_more else None,
        },
    )


@router.post("/pdf-split")
async def split_pdf(payload: PdfSplitRequest, request: Request, account: dict = Depends(require_csrf)):
    """Split one verified private Asset Vault PDF into a private sanitized PDF."""
    _require_enabled()
    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copy: Path | None = None
    final_path: Path | None = None
    source_asset_id = payload.source_asset_id

    ensure_copyfast_schema()
    with transaction() as conn:
        source_row = conn.execute(
            """SELECT id, project_id, extension, content_type, byte_size, sha256, storage_key, state
               FROM web_asset_files WHERE id=? AND account_id=?""",
            (source_asset_id, account_id),
        ).fetchone()
        if not source_row or str(source_row[7]) != "active":
            return _source_not_found()
        if str(source_row[2]) != ".pdf" or str(source_row[3]) != "application/pdf":
            raise HTTPException(status_code=422, detail="PDF Split chỉ nhận PDF private hợp lệ trong Asset Vault")
        source_bytes = int(source_row[4])
        if source_bytes < 1 or source_bytes > MAX_INPUT_BYTES:
            raise HTTPException(status_code=413, detail="PDF nguồn vượt giới hạn 20 MB")
        source_sha256 = str(source_row[5])
        request_fingerprint = _request_fingerprint(
            source_asset_id=source_asset_id,
            page_range=payload.page_range,
            source_sha256=source_sha256,
            source_bytes=source_bytes,
        )
        existing = conn.execute(
            f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                WHERE account_id=? AND kind=? AND idempotency_key=?""",
            (account_id, PDF_SPLIT_KIND, payload.idempotency_key),
        ).fetchone()
        if existing:
            if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                raise HTTPException(status_code=409, detail="Idempotency key đã được dùng cho PDF Split khác")
            return _operation_response(_operation_public(tuple(existing[:-1])))

        operation_id = str(uuid.uuid4())
        now = utc_now()
        conn.execute(
            """INSERT INTO web_document_operations
               (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                created_at, queued_at, started_at, updated_at)
               VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                operation_id, account_id, source_asset_id, str(source_row[1]) if source_row[1] else None,
                PDF_SPLIT_KIND, payload.idempotency_key, request_fingerprint, source_sha256, source_bytes,
                1, payload.page_range, now, now, now, now,
            ),
        )
        _record_event(conn, operation_id=operation_id, state="queued", when=now)
        conn.execute(
            "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
            (now, operation_id, account_id),
        )
        _record_event(conn, operation_id=operation_id, state="processing", when=now)
        source_storage_key = str(source_row[6])

    try:
        source_path = _asset_path(asset_vault_directory(), source_storage_key)
        source_copy = _staging_path(root, ".source.pdf")
        _copy_verified_source(source_path, source_copy, expected_bytes=source_bytes, expected_digest=source_sha256)
        final_path, output_storage_key, output_bytes, output_digest, source_page_count, start_page, end_page = _build_split_output(
            root,
            source_copy,
            page_range=payload.page_range,
        )
        output_pages = end_page - start_page + 1
        filename = f"toan-aas-pdf-pages-{start_page}-{end_page}.pdf"
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, PDF_SPLIT_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("PDF Split không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', selected_start_page=?, selected_end_page=?, source_page_count=?,
                       output_page_count=?, storage_key=?, original_filename=?, content_type='application/pdf',
                       byte_size=?, sha256=?, completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    start_page, end_page, source_page_count, output_pages, output_storage_key, filename,
                    output_bytes, output_digest, now, now, operation_id, account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.pdf_split",
                request_id=_request_id(request),
                target=operation_id,
                detail=f"pages={start_page}-{end_page};source_pages={source_page_count};bytes={output_bytes}",
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc PDF Split vừa hoàn tất")
        final_path = None  # Output now has committed metadata ownership.
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code == "PDF_SOURCE_UNAVAILABLE":
            _mark_source_unavailable(source_asset_id, account_id)
        _mark_failed(operation_id, account_id, kind=PDF_SPLIT_KIND, request=request, code=exc.code)
        raise HTTPException(status_code=422, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=PDF_SPLIT_KIND, request=request, code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION")
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=PDF_SPLIT_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể tách PDF an toàn") from exc
    finally:
        _safe_unlink(source_copy)


def _merge_sources_for_account(conn, *, source_asset_ids: list[str], account_id: str) -> tuple[list[dict[str, Any]], str | None]:
    """Load ordered, owner-scoped PDF sources and apply aggregate bounds."""
    sources: list[dict[str, Any]] = []
    total_bytes = 0
    project_ids: set[str] = set()
    for asset_id in source_asset_ids:
        row = conn.execute(
            """SELECT id, project_id, extension, content_type, byte_size, sha256, storage_key, state
               FROM web_asset_files WHERE id=? AND account_id=?""",
            (asset_id, account_id),
        ).fetchone()
        if not row or str(row[7]) != "active":
            raise DocumentOperationError("Không tìm thấy PDF private đang hoạt động thuộc Web account hiện tại.", code="PDF_SOURCE_NOT_FOUND")
        if str(row[2]) != ".pdf" or str(row[3]) != "application/pdf":
            raise DocumentOperationError("PDF Merge chỉ nhận PDF private hợp lệ trong Asset Vault", code="PDF_SOURCE_INVALID")
        byte_size = int(row[4])
        if byte_size < 1 or byte_size > MAX_INPUT_BYTES:
            raise DocumentOperationError("Mỗi PDF nguồn không được vượt quá 20 MB", code="PDF_INPUT_TOO_LARGE")
        total_bytes += byte_size
        if total_bytes > MAX_MERGE_INPUT_BYTES:
            raise DocumentOperationError("Tổng dung lượng PDF nguồn vượt giới hạn 40 MB", code="PDF_MERGE_INPUT_LIMIT")
        source_sha256 = str(row[5] or "")
        if not re.fullmatch(r"[0-9a-f]{64}", source_sha256):
            raise DocumentOperationError("PDF nguồn không còn sẵn sàng", code="PDF_SOURCE_UNAVAILABLE")
        storage_key = str(row[6] or "")
        if not ASSET_STORAGE_KEY_PATTERN.fullmatch(storage_key):
            raise DocumentOperationError("PDF nguồn không còn sẵn sàng", code="PDF_SOURCE_UNAVAILABLE")
        project_id = str(row[1]) if row[1] else None
        if project_id:
            project_ids.add(project_id)
        sources.append(
            {
                "id": str(row[0]),
                "project_id": project_id,
                "byte_size": byte_size,
                "sha256": source_sha256,
                "storage_key": storage_key,
            }
        )
    if len(sources) < 2:
        raise DocumentOperationError("Cần ít nhất hai PDF private để gộp", code="PDF_MERGE_SOURCE_COUNT")
    project_id = next(iter(project_ids)) if len(project_ids) == 1 else None
    return sources, project_id


def _image_sources_for_account(conn, *, source_asset_ids: list[str], account_id: str) -> tuple[list[dict[str, Any]], str | None]:
    """Load ordered, owner-scoped static image sources and enforce batch bounds."""
    sources: list[dict[str, Any]] = []
    total_bytes = 0
    project_ids: set[str] = set()
    for asset_id in source_asset_ids:
        row = conn.execute(
            """SELECT id, project_id, extension, content_type, byte_size, sha256, storage_key, state
               FROM web_asset_files WHERE id=? AND account_id=?""",
            (asset_id, account_id),
        ).fetchone()
        if not row or str(row[7]) != "active":
            raise DocumentOperationError("Không tìm thấy ảnh private đang hoạt động thuộc Web account hiện tại.", code="IMAGE_SOURCE_NOT_FOUND")
        extension = str(row[2] or "").lower()
        content_type = str(row[3] or "").lower()
        if IMAGE_INPUT_MIME_BY_EXTENSION.get(extension) != content_type:
            raise DocumentOperationError("Ảnh → PDF chỉ nhận JPEG, PNG hoặc WebP private hợp lệ trong Asset Vault", code="IMAGE_SOURCE_INVALID")
        byte_size = int(row[4])
        if byte_size < 1 or byte_size > MAX_INPUT_BYTES:
            raise DocumentOperationError("Mỗi ảnh nguồn không được vượt quá 20 MB", code="IMAGE_INPUT_TOO_LARGE")
        total_bytes += byte_size
        if total_bytes > MAX_IMAGE_PDF_INPUT_BYTES:
            raise DocumentOperationError("Tổng dung lượng ảnh nguồn vượt giới hạn 40 MB", code="IMAGE_TO_PDF_INPUT_LIMIT")
        source_sha256 = str(row[5] or "")
        if not re.fullmatch(r"[0-9a-f]{64}", source_sha256):
            raise DocumentOperationError("Ảnh nguồn không còn sẵn sàng", code="IMAGE_SOURCE_UNAVAILABLE")
        storage_key = str(row[6] or "")
        if not ASSET_STORAGE_KEY_PATTERN.fullmatch(storage_key):
            raise DocumentOperationError("Ảnh nguồn không còn sẵn sàng", code="IMAGE_SOURCE_UNAVAILABLE")
        project_id = str(row[1]) if row[1] else None
        if project_id:
            project_ids.add(project_id)
        sources.append(
            {
                "id": str(row[0]),
                "project_id": project_id,
                "extension": extension,
                "byte_size": byte_size,
                "sha256": source_sha256,
                "storage_key": storage_key,
            }
        )
    if not 1 <= len(sources) <= MAX_IMAGE_PDF_SOURCES:
        raise DocumentOperationError("Cần từ 1 đến 8 ảnh private để tạo PDF", code="IMAGE_SOURCE_COUNT")
    project_id = next(iter(project_ids)) if len(project_ids) == 1 else None
    return sources, project_id


@router.post("/pdf-merge")
async def merge_pdf(payload: PdfMergeRequest, request: Request, account: dict = Depends(require_csrf)):
    """Merge ordered verified Asset Vault PDFs into one sanitized private PDF."""
    _require_enabled()
    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copies: list[Path] = []
    final_path: Path | None = None
    sources: list[dict[str, Any]] = []
    active_source_id: str | None = None

    ensure_copyfast_schema()
    with transaction() as conn:
        try:
            sources, scoped_project_id = _merge_sources_for_account(
                conn,
                source_asset_ids=payload.source_asset_ids,
                account_id=account_id,
            )
        except DocumentOperationError as exc:
            if exc.code == "PDF_SOURCE_NOT_FOUND":
                return _source_not_found()
            status_code = 413 if exc.code in {"PDF_INPUT_TOO_LARGE", "PDF_MERGE_INPUT_LIMIT"} else 422
            raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
        request_fingerprint = _merge_request_fingerprint(sources)
        existing = conn.execute(
            f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                WHERE account_id=? AND kind=? AND idempotency_key=?""",
            (account_id, PDF_MERGE_KIND, payload.idempotency_key),
        ).fetchone()
        if existing:
            if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                raise HTTPException(status_code=409, detail="Idempotency key đã được dùng cho PDF Merge khác")
            return _operation_response(_operation_public(tuple(existing[:-1])))

        operation_id = str(uuid.uuid4())
        now = utc_now()
        first = sources[0]
        conn.execute(
            """INSERT INTO web_document_operations
               (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                created_at, queued_at, started_at, updated_at)
               VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, ?, '', ?, ?, ?, ?)""",
            (
                operation_id, account_id, first["id"], scoped_project_id, PDF_MERGE_KIND,
                payload.idempotency_key, request_fingerprint, first["sha256"], first["byte_size"],
                len(sources), now, now, now, now,
            ),
        )
        for source_index, source in enumerate(sources, start=1):
            conn.execute(
                """INSERT INTO web_document_operation_sources
                   (id, operation_id, source_asset_id, source_index, source_sha256, source_byte_size, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (
                    str(uuid.uuid4()), operation_id, source["id"], source_index,
                    source["sha256"], source["byte_size"], now,
                ),
            )
        _record_event(conn, operation_id=operation_id, state="queued", when=now)
        conn.execute(
            "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
            (now, operation_id, account_id),
        )
        _record_event(conn, operation_id=operation_id, state="processing", when=now)

    try:
        asset_root = asset_vault_directory()
        for source in sources:
            active_source_id = str(source["id"])
            source_path = _asset_path(asset_root, source["storage_key"])
            source_copy = _staging_path(root, ".source.pdf")
            # Track the private staging name before any I/O so a failed
            # integrity copy cannot leave an unreferenced source behind.
            source_copies.append(source_copy)
            _copy_verified_source(
                source_path,
                source_copy,
                expected_bytes=int(source["byte_size"]),
                expected_digest=str(source["sha256"]),
            )
        final_path, output_storage_key, output_bytes, output_digest, source_page_count = _build_merge_output(root, source_copies)
        filename = "toan-aas-merged-pdf.pdf"
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, PDF_MERGE_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("PDF Merge không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', source_page_count=?, output_page_count=?, storage_key=?,
                       original_filename=?, content_type='application/pdf', byte_size=?, sha256=?,
                       completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    source_page_count, source_page_count, output_storage_key, filename, output_bytes,
                    output_digest, now, now, operation_id, account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.pdf_merge",
                request_id=_request_id(request),
                target=operation_id,
                detail=f"sources={len(sources)};pages={source_page_count};bytes={output_bytes}",
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc PDF Merge vừa hoàn tất")
        final_path = None
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code == "PDF_SOURCE_UNAVAILABLE" and active_source_id:
            _mark_source_unavailable(active_source_id, account_id)
        _mark_failed(operation_id, account_id, kind=PDF_MERGE_KIND, request=request, code=exc.code)
        raise HTTPException(status_code=422, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(
            operation_id,
            account_id,
            kind=PDF_MERGE_KIND,
            request=request,
            code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION",
        )
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=PDF_MERGE_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể gộp PDF an toàn") from exc
    finally:
        for source_copy in source_copies:
            _safe_unlink(source_copy)


@router.post("/image-to-pdf")
async def image_to_pdf(
    payload: ImageToPdfRequest,
    request: Request,
    account: dict = Depends(require_csrf),
):
    """Create a verified one-page-per-image PDF from ordered private Vault sources."""
    _require_image_to_pdf_enabled()
    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copies: list[Path] = []
    final_path: Path | None = None
    sources: list[dict[str, Any]] = []
    active_source_id: str | None = None
    capacity_reserved = False

    ensure_copyfast_schema()
    try:
        with transaction() as conn:
            try:
                sources, scoped_project_id = _image_sources_for_account(
                    conn,
                    source_asset_ids=payload.source_asset_ids,
                    account_id=account_id,
                )
            except DocumentOperationError as exc:
                if exc.code == "IMAGE_SOURCE_NOT_FOUND":
                    return _image_source_not_found()
                status_code = 413 if exc.code in {"IMAGE_INPUT_TOO_LARGE", "IMAGE_TO_PDF_INPUT_LIMIT"} else 422
                raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
            request_fingerprint = _image_to_pdf_request_fingerprint(sources)
            existing = conn.execute(
                f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                    WHERE account_id=? AND kind=? AND idempotency_key=?""",
                (account_id, IMAGE_TO_PDF_KIND, payload.idempotency_key),
            ).fetchone()
            if existing:
                if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                    raise HTTPException(status_code=409, detail="Idempotency key đã được dùng cho Ảnh → PDF khác")
                return _operation_response(_operation_public(tuple(existing[:-1])))

            # The request is new only after its authenticated, owner-scoped
            # idempotency lookup. This keeps a completed retry available even
            # while another decoder-heavy batch is consuming the one slot.
            _reserve_image_to_pdf_capacity()
            capacity_reserved = True
            operation_id = str(uuid.uuid4())
            now = utc_now()
            first = sources[0]
            conn.execute(
                """INSERT INTO web_document_operations
                   (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                    request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                    created_at, queued_at, started_at, updated_at)
                   VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, ?, '', ?, ?, ?, ?)""",
                (
                    operation_id, account_id, first["id"], scoped_project_id, IMAGE_TO_PDF_KIND,
                    payload.idempotency_key, request_fingerprint, first["sha256"], first["byte_size"],
                    len(sources), now, now, now, now,
                ),
            )
            for source_index, source in enumerate(sources, start=1):
                conn.execute(
                    """INSERT INTO web_document_operation_sources
                       (id, operation_id, source_asset_id, source_index, source_sha256, source_byte_size, created_at)
                       VALUES (?, ?, ?, ?, ?, ?, ?)""",
                    (
                        str(uuid.uuid4()), operation_id, source["id"], source_index,
                        source["sha256"], source["byte_size"], now,
                    ),
                )
            _record_event(conn, operation_id=operation_id, state="queued", when=now)
            conn.execute(
                "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
                (now, operation_id, account_id),
            )
            _record_event(conn, operation_id=operation_id, state="processing", when=now)
    except Exception:
        if capacity_reserved:
            _IMAGE_TO_PDF_CAPACITY.release()
        raise

    try:
        asset_root = asset_vault_directory()
        for source in sources:
            active_source_id = str(source["id"])
            source_path = _asset_path(asset_root, str(source["storage_key"]))
            source_copy = _staging_path(root, f".source{source['extension']}")
            # Track each staging name before I/O so partial copy failures are
            # always removed from the private processing directory.
            source_copies.append(source_copy)
            await run_in_threadpool(
                _copy_verified_image_source,
                source_path,
                source_copy,
                extension=str(source["extension"]),
                expected_bytes=int(source["byte_size"]),
                expected_digest=str(source["sha256"]),
            )
        final_path, output_storage_key, output_bytes, output_digest, output_pages, total_pixels = await run_in_threadpool(
            _build_image_to_pdf_output,
            root,
            [(source_copy, str(source["extension"])) for source_copy, source in zip(source_copies, sources, strict=True)],
        )
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, IMAGE_TO_PDF_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("Ảnh → PDF không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', source_page_count=?, output_page_count=?, storage_key=?,
                       original_filename='toan-aas-images.pdf', content_type='application/pdf',
                       byte_size=?, sha256=?, completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    output_pages, output_pages, output_storage_key, output_bytes, output_digest,
                    now, now, operation_id, account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.image_to_pdf",
                request_id=_request_id(request),
                target=operation_id,
                detail=f"sources={len(sources)};pixels={total_pixels};bytes={output_bytes}",
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc Ảnh → PDF vừa hoàn tất")
        final_path = None
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code == "IMAGE_SOURCE_UNAVAILABLE" and active_source_id:
            _mark_source_unavailable(active_source_id, account_id)
        _mark_failed(operation_id, account_id, kind=IMAGE_TO_PDF_KIND, request=request, code=exc.code)
        status_code = 413 if exc.code in {
            "IMAGE_INPUT_TOO_LARGE",
            "IMAGE_TO_PDF_INPUT_LIMIT",
            "IMAGE_DIMENSION_LIMIT",
            "IMAGE_ASPECT_RATIO_LIMIT",
            "IMAGE_PIXEL_LIMIT",
            "IMAGE_TOTAL_PIXEL_LIMIT",
            "IMAGE_RESOURCE_LIMIT",
            "PDF_OUTPUT_LIMIT",
        } else 422
        raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(
            operation_id,
            account_id,
            kind=IMAGE_TO_PDF_KIND,
            request=request,
            code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION",
        )
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=IMAGE_TO_PDF_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể tạo PDF từ ảnh an toàn") from exc
    finally:
        for source_copy in source_copies:
            _safe_unlink(source_copy)
        if capacity_reserved:
            _IMAGE_TO_PDF_CAPACITY.release()


@router.post("/ocr-image")
async def ocr_image(payload: ImageOcrRequest, request: Request, account: dict = Depends(require_csrf)):
    """Extract text from one private image with local, opt-in OCR only.

    The server derives replay protection from the immutable Asset Vault source
    revision.  No request field can select a provider, upload bytes, reference
    a URL/path or pass a browser idempotency token into this execution path.
    """

    _require_image_ocr_enabled()
    try:
        ocr_module, installed_languages = _image_ocr_runtime()
        ocr_language = _image_ocr_language(payload.language, installed_languages)
    except DocumentOperationError as exc:
        return _ocr_readiness_guard(exc)

    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copy: Path | None = None
    final_path: Path | None = None
    source_asset_id = payload.source_asset_id
    source: dict[str, Any] | None = None
    capacity_reserved = False

    ensure_copyfast_schema()
    try:
        with transaction() as conn:
            try:
                sources, scoped_project_id = _image_sources_for_account(
                    conn,
                    source_asset_ids=[source_asset_id],
                    account_id=account_id,
                )
            except DocumentOperationError as exc:
                if exc.code == "IMAGE_SOURCE_NOT_FOUND":
                    return _ocr_source_not_found()
                status_code = 413 if exc.code in {
                    "IMAGE_INPUT_TOO_LARGE",
                    "IMAGE_TO_PDF_INPUT_LIMIT",
                } else 422
                raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
            source = sources[0]
            request_fingerprint = _image_ocr_request_fingerprint(
                source_asset_id=source_asset_id,
                source_sha256=str(source["sha256"]),
                source_bytes=int(source["byte_size"]),
                language=payload.language,
            )
            # The client never controls this key. It only identifies the same
            # owner-scoped source revision + language selection, so refreshes
            # cannot start concurrent duplicate OCR work or artifacts.
            server_idempotency_key = f"ocr-{request_fingerprint}"
            existing = conn.execute(
                f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                    WHERE account_id=? AND kind=? AND idempotency_key=?""",
                (account_id, IMAGE_OCR_KIND, server_idempotency_key),
            ).fetchone()
            if existing:
                if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                    raise RuntimeError("Idempotency OCR không khớp fingerprint nguồn")
                existing_operation = _operation_public(tuple(existing[:-1]))
                # A true empty-text result is a final truthful outcome for the
                # same immutable source.  A local runtime/language failure is
                # transient instead: once readiness above succeeds, allow the
                # same derived idempotency record to resume rather than making
                # the owner create a changed source merely to retry.
                if not (
                    existing_operation.get("state") == "guarded"
                    and str(existing_operation.get("_failure_code") or "")
                    in {"OCR_RUNTIME_UNAVAILABLE", "OCR_LANGUAGE_UNAVAILABLE"}
                ):
                    return _operation_response(existing_operation)
                _reserve_image_ocr_capacity()
                capacity_reserved = True
                operation_id = str(existing_operation["id"])
                now = utc_now()
                conn.execute(
                    """UPDATE web_document_operations
                       SET state='processing', failure_code=NULL, storage_key=NULL, original_filename=NULL,
                           content_type=NULL, byte_size=NULL, sha256=NULL, completed_at=NULL,
                           started_at=?, updated_at=?
                       WHERE id=? AND account_id=? AND kind=? AND state='guarded'""",
                    (now, now, operation_id, account_id, IMAGE_OCR_KIND),
                )
                _record_event(conn, operation_id=operation_id, state="processing", when=now)
            else:
                _reserve_image_ocr_capacity()
                capacity_reserved = True
                operation_id = str(uuid.uuid4())
                now = utc_now()
                conn.execute(
                    """INSERT INTO web_document_operations
                       (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                        request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                        created_at, queued_at, started_at, updated_at)
                       VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, 1, ?, ?, ?, ?, ?)""",
                    (
                        operation_id,
                        account_id,
                        source_asset_id,
                        scoped_project_id,
                        IMAGE_OCR_KIND,
                        server_idempotency_key,
                        request_fingerprint,
                        str(source["sha256"]),
                        int(source["byte_size"]),
                        f"ocr:{payload.language}",
                        now,
                        now,
                        now,
                        now,
                    ),
                )
                _record_event(conn, operation_id=operation_id, state="queued", when=now)
                conn.execute(
                    "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
                    (now, operation_id, account_id),
                )
                _record_event(conn, operation_id=operation_id, state="processing", when=now)
    except Exception:
        if capacity_reserved:
            _IMAGE_OCR_CAPACITY.release()
        raise

    try:
        if source is None:
            raise RuntimeError("OCR image source không còn sẵn sàng")
        source_path = _asset_path(asset_vault_directory(), str(source["storage_key"]))
        source_copy = _staging_path(root, f".source{source['extension']}")
        await run_in_threadpool(
            _copy_verified_image_source,
            source_path,
            source_copy,
            extension=str(source["extension"]),
            expected_bytes=int(source["byte_size"]),
            expected_digest=str(source["sha256"]),
        )
        (
            final_path,
            output_storage_key,
            output_bytes,
            output_digest,
            pixels,
            characters,
        ) = await run_in_threadpool(
            _build_image_ocr_output,
            root,
            source_copy,
            str(source["extension"]),
            language=ocr_language,
            ocr_module=ocr_module,
        )
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, IMAGE_OCR_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("OCR ảnh không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', source_page_count=1, output_page_count=NULL, storage_key=?,
                       original_filename='toan-aas-image-ocr.txt', content_type='text/plain; charset=utf-8',
                       byte_size=?, sha256=?, completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    output_storage_key,
                    output_bytes,
                    output_digest,
                    now,
                    now,
                    operation_id,
                    account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.image_ocr",
                request_id=_request_id(request),
                target=operation_id,
                detail=(
                    f"language={payload.language};pixels={pixels};characters={characters};bytes={output_bytes}"
                ),
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc OCR ảnh vừa hoàn tất")
        final_path = None
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code in {"OCR_TEXT_NOT_FOUND", "OCR_RUNTIME_UNAVAILABLE", "OCR_LANGUAGE_UNAVAILABLE"}:
            _mark_guarded(operation_id, account_id, kind=IMAGE_OCR_KIND, request=request, code=exc.code)
            with transaction() as conn:
                guarded = conn.execute(
                    f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                    (operation_id, account_id),
                ).fetchone()
            if guarded:
                return _operation_response(_operation_public(tuple(guarded)))
        if exc.code == "IMAGE_SOURCE_UNAVAILABLE":
            _mark_source_unavailable(source_asset_id, account_id)
        _mark_failed(operation_id, account_id, kind=IMAGE_OCR_KIND, request=request, code=exc.code)
        status_code = 413 if exc.code in {
            "IMAGE_INPUT_TOO_LARGE",
            "IMAGE_DIMENSION_LIMIT",
            "IMAGE_ASPECT_RATIO_LIMIT",
            "IMAGE_PIXEL_LIMIT",
            "IMAGE_RESOURCE_LIMIT",
            "OCR_OUTPUT_LIMIT",
        } else 422
        raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(
            operation_id,
            account_id,
            kind=IMAGE_OCR_KIND,
            request=request,
            code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION",
        )
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=IMAGE_OCR_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể OCR ảnh an toàn") from exc
    finally:
        _safe_unlink(source_copy)
        if capacity_reserved:
            _IMAGE_OCR_CAPACITY.release()


@router.post("/pdf-to-images")
async def pdf_to_images(payload: PdfToImagesRequest, request: Request, account: dict = Depends(require_csrf)):
    """Render a verified private PDF to a checked PNG or deterministic PNG ZIP."""
    _require_pdf_to_images_enabled()
    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copy: Path | None = None
    final_path: Path | None = None
    source_asset_id = payload.source_asset_id
    source_storage_key = ""
    source_bytes = 0
    source_sha256 = ""
    capacity_reserved = False

    ensure_copyfast_schema()
    try:
        with transaction() as conn:
            source_row = conn.execute(
                """SELECT id, project_id, extension, content_type, byte_size, sha256, storage_key, state
                   FROM web_asset_files WHERE id=? AND account_id=?""",
                (source_asset_id, account_id),
            ).fetchone()
            if not source_row or str(source_row[7]) != "active":
                return _source_not_found()
            if str(source_row[2]) != ".pdf" or str(source_row[3]) != "application/pdf":
                raise HTTPException(status_code=422, detail="PDF → ảnh chỉ nhận PDF private hợp lệ trong Asset Vault")
            source_bytes = int(source_row[4])
            if source_bytes < 1 or source_bytes > MAX_INPUT_BYTES:
                raise HTTPException(status_code=413, detail="PDF nguồn vượt giới hạn 20 MB")
            source_sha256 = str(source_row[5] or "")
            source_storage_key = str(source_row[6] or "")
            if not re.fullmatch(r"[0-9a-f]{64}", source_sha256) or not ASSET_STORAGE_KEY_PATTERN.fullmatch(source_storage_key):
                raise HTTPException(status_code=422, detail="PDF nguồn không còn sẵn sàng")
            request_fingerprint = _pdf_to_images_request_fingerprint(
                source_asset_id=source_asset_id,
                source_sha256=source_sha256,
                source_bytes=source_bytes,
            )
            existing = conn.execute(
                f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                    WHERE account_id=? AND kind=? AND idempotency_key=?""",
                (account_id, PDF_TO_IMAGES_KIND, payload.idempotency_key),
            ).fetchone()
            if existing:
                if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                    raise HTTPException(status_code=409, detail="Idempotency key đã được dùng cho PDF → ảnh khác")
                return _operation_response(_operation_public(tuple(existing[:-1])))

            # Replays are resolved above before capacity is acquired. A new
            # work item must have a render slot before it gets a lifecycle row.
            _reserve_pdf_to_images_capacity()
            capacity_reserved = True
            operation_id = str(uuid.uuid4())
            now = utc_now()
            conn.execute(
                """INSERT INTO web_document_operations
                   (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                    request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                    created_at, queued_at, started_at, updated_at)
                   VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, 1, '', ?, ?, ?, ?)""",
                (
                    operation_id,
                    account_id,
                    source_asset_id,
                    str(source_row[1]) if source_row[1] else None,
                    PDF_TO_IMAGES_KIND,
                    payload.idempotency_key,
                    request_fingerprint,
                    source_sha256,
                    source_bytes,
                    now,
                    now,
                    now,
                    now,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="queued", when=now)
            conn.execute(
                "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
                (now, operation_id, account_id),
            )
            _record_event(conn, operation_id=operation_id, state="processing", when=now)
    except Exception:
        if capacity_reserved:
            _PDF_TO_IMAGES_CAPACITY.release()
        raise

    try:
        source_path = _asset_path(asset_vault_directory(), source_storage_key)
        source_copy = _staging_path(root, ".source.pdf")
        await run_in_threadpool(
            _copy_verified_source,
            source_path,
            source_copy,
            expected_bytes=source_bytes,
            expected_digest=source_sha256,
        )
        (
            final_path,
            output_storage_key,
            output_bytes,
            output_digest,
            source_page_count,
            output_page_count,
            media_type,
            output_filename,
        ) = await run_in_threadpool(_build_pdf_to_images_output, root, source_copy)
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, PDF_TO_IMAGES_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("PDF → ảnh không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', source_page_count=?, output_page_count=?, storage_key=?,
                       original_filename=?, content_type=?, byte_size=?, sha256=?,
                       completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    source_page_count,
                    output_page_count,
                    output_storage_key,
                    output_filename,
                    media_type,
                    output_bytes,
                    output_digest,
                    now,
                    now,
                    operation_id,
                    account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.pdf_to_images",
                request_id=_request_id(request),
                target=operation_id,
                detail=(
                    f"source_pages={source_page_count};output_pages={output_page_count};"
                    f"artifact={'png' if output_page_count == 1 else 'zip'};bytes={output_bytes}"
                ),
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc PDF → ảnh vừa hoàn tất")
        final_path = None
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code == "PDF_SOURCE_UNAVAILABLE":
            _mark_source_unavailable(source_asset_id, account_id)
        _mark_failed(operation_id, account_id, kind=PDF_TO_IMAGES_KIND, request=request, code=exc.code)
        status_code = 413 if exc.code in {
            "PDF_INPUT_TOO_LARGE",
            "PDF_PAGE_LIMIT",
            "PDF_RENDER_DIMENSION_LIMIT",
            "PDF_RENDER_PIXEL_LIMIT",
            "PDF_RENDER_TOTAL_PIXEL_LIMIT",
            "PDF_RENDER_RESOURCE_LIMIT",
            "PDF_TO_IMAGES_OUTPUT_LIMIT",
        } else 422
        raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(
            operation_id,
            account_id,
            kind=PDF_TO_IMAGES_KIND,
            request=request,
            code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION",
        )
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=PDF_TO_IMAGES_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể render PDF sang ảnh an toàn") from exc
    finally:
        _safe_unlink(source_copy)
        if capacity_reserved:
            _PDF_TO_IMAGES_CAPACITY.release()


@router.post("/pdf-to-word")
async def pdf_to_word(payload: PdfToWordRequest, request: Request, account: dict = Depends(require_csrf)):
    """Export only real selectable PDF text into a private verified DOCX."""
    _require_pdf_to_word_enabled()
    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copy: Path | None = None
    final_path: Path | None = None
    source_asset_id = payload.source_asset_id
    source_storage_key = ""
    source_bytes = 0
    source_sha256 = ""
    capacity_reserved = False

    ensure_copyfast_schema()
    try:
        with transaction() as conn:
            source_row = conn.execute(
                """SELECT id, project_id, extension, content_type, byte_size, sha256, storage_key, state
                   FROM web_asset_files WHERE id=? AND account_id=?""",
                (source_asset_id, account_id),
            ).fetchone()
            if not source_row or str(source_row[7]) != "active":
                return _source_not_found()
            if str(source_row[2]) != ".pdf" or str(source_row[3]) != "application/pdf":
                raise HTTPException(status_code=422, detail="PDF có text → Word chỉ nhận PDF private hợp lệ trong Asset Vault")
            source_bytes = int(source_row[4])
            if source_bytes < 1 or source_bytes > MAX_INPUT_BYTES:
                raise HTTPException(status_code=413, detail="PDF nguồn vượt giới hạn 20 MB")
            source_sha256 = str(source_row[5] or "")
            source_storage_key = str(source_row[6] or "")
            if not re.fullmatch(r"[0-9a-f]{64}", source_sha256) or not ASSET_STORAGE_KEY_PATTERN.fullmatch(source_storage_key):
                raise HTTPException(status_code=422, detail="PDF nguồn không còn sẵn sàng")
            request_fingerprint = _pdf_to_word_request_fingerprint(
                source_asset_id=source_asset_id,
                source_sha256=source_sha256,
                source_bytes=source_bytes,
            )
            existing = conn.execute(
                f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                    WHERE account_id=? AND kind=? AND idempotency_key=?""",
                (account_id, PDF_TO_WORD_KIND, payload.idempotency_key),
            ).fetchone()
            if existing:
                if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                    raise HTTPException(status_code=409, detail="Idempotency key đã được dùng cho PDF có text → Word khác")
                return _operation_response(_operation_public(tuple(existing[:-1])))

            # A replayed completed/guarded request must remain readable even
            # while the one extraction slot is occupied. Only a truly new job
            # reserves capacity and creates a lifecycle row.
            _reserve_pdf_to_word_capacity()
            capacity_reserved = True
            operation_id = str(uuid.uuid4())
            now = utc_now()
            conn.execute(
                """INSERT INTO web_document_operations
                   (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                    request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                    created_at, queued_at, started_at, updated_at)
                   VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, 1, '', ?, ?, ?, ?)""",
                (
                    operation_id,
                    account_id,
                    source_asset_id,
                    str(source_row[1]) if source_row[1] else None,
                    PDF_TO_WORD_KIND,
                    payload.idempotency_key,
                    request_fingerprint,
                    source_sha256,
                    source_bytes,
                    now,
                    now,
                    now,
                    now,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="queued", when=now)
            conn.execute(
                "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
                (now, operation_id, account_id),
            )
            _record_event(conn, operation_id=operation_id, state="processing", when=now)
    except Exception:
        if capacity_reserved:
            _PDF_TO_WORD_CAPACITY.release()
        raise

    try:
        source_path = _asset_path(asset_vault_directory(), source_storage_key)
        source_copy = _staging_path(root, ".source.pdf")
        await run_in_threadpool(
            _copy_verified_source,
            source_path,
            source_copy,
            expected_bytes=source_bytes,
            expected_digest=source_sha256,
        )
        (
            final_path,
            output_storage_key,
            output_bytes,
            output_digest,
            source_page_count,
            extracted_characters,
            paragraph_count,
        ) = await run_in_threadpool(_build_pdf_to_word_output, root, source_copy)
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, PDF_TO_WORD_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("PDF có text → Word không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', source_page_count=?, output_page_count=NULL, storage_key=?,
                       original_filename='toan-aas-pdf-text.docx', content_type=?,
                       byte_size=?, sha256=?, completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    source_page_count,
                    output_storage_key,
                    DOCX_MEDIA_TYPE,
                    output_bytes,
                    output_digest,
                    now,
                    now,
                    operation_id,
                    account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.pdf_to_word_text",
                request_id=_request_id(request),
                target=operation_id,
                detail=(
                    f"source_pages={source_page_count};characters={extracted_characters};"
                    f"paragraphs={paragraph_count};bytes={output_bytes}"
                ),
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc PDF có text → Word vừa hoàn tất")
        final_path = None
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code == "PDF_TEXT_NOT_FOUND":
            _mark_guarded(operation_id, account_id, kind=PDF_TO_WORD_KIND, request=request, code=exc.code)
        else:
            if exc.code == "PDF_SOURCE_UNAVAILABLE":
                _mark_source_unavailable(source_asset_id, account_id)
            _mark_failed(operation_id, account_id, kind=PDF_TO_WORD_KIND, request=request, code=exc.code)
        status_code = 413 if exc.code in {
            "PDF_INPUT_TOO_LARGE",
            "PDF_PAGE_LIMIT",
            "PDF_TEXT_LIMIT",
            "PDF_TO_WORD_OUTPUT_LIMIT",
        } else 422
        raise HTTPException(status_code=status_code, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(
            operation_id,
            account_id,
            kind=PDF_TO_WORD_KIND,
            request=request,
            code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION",
        )
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=PDF_TO_WORD_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể xuất DOCX từ PDF an toàn") from exc
    finally:
        _safe_unlink(source_copy)
        if capacity_reserved:
            _PDF_TO_WORD_CAPACITY.release()


@router.post("/pdf-optimize")
async def optimize_pdf(payload: PdfOptimizeRequest, request: Request, account: dict = Depends(require_csrf)):
    """Optimize one verified Asset Vault PDF only if it becomes smaller."""
    _require_enabled()
    root = document_operations_directory()
    account_id = str(account["id"])
    operation_id = ""
    source_copy: Path | None = None
    final_path: Path | None = None
    source_asset_id = payload.source_asset_id

    ensure_copyfast_schema()
    with transaction() as conn:
        source_row = conn.execute(
            """SELECT id, project_id, extension, content_type, byte_size, sha256, storage_key, state
               FROM web_asset_files WHERE id=? AND account_id=?""",
            (source_asset_id, account_id),
        ).fetchone()
        if not source_row or str(source_row[7]) != "active":
            return _source_not_found()
        if str(source_row[2]) != ".pdf" or str(source_row[3]) != "application/pdf":
            raise HTTPException(status_code=422, detail="PDF Optimize chỉ nhận PDF private hợp lệ trong Asset Vault")
        source_bytes = int(source_row[4])
        if source_bytes < 1 or source_bytes > MAX_INPUT_BYTES:
            raise HTTPException(status_code=413, detail="PDF nguồn vượt giới hạn 20 MB")
        source_sha256 = str(source_row[5] or "")
        source_storage_key = str(source_row[6] or "")
        if not re.fullmatch(r"[0-9a-f]{64}", source_sha256) or not ASSET_STORAGE_KEY_PATTERN.fullmatch(source_storage_key):
            raise HTTPException(status_code=422, detail="PDF nguồn không còn sẵn sàng")
        request_fingerprint = _optimize_request_fingerprint(
            source_asset_id=source_asset_id,
            source_sha256=source_sha256,
            source_bytes=source_bytes,
        )
        existing = conn.execute(
            f"""SELECT {OPERATION_SELECT}, request_fingerprint FROM web_document_operations
                WHERE account_id=? AND kind=? AND idempotency_key=?""",
            (account_id, PDF_OPTIMIZE_KIND, payload.idempotency_key),
        ).fetchone()
        if existing:
            if not hmac.compare_digest(str(existing[-1] or ""), request_fingerprint):
                raise HTTPException(status_code=409, detail="Idempotency key đã được dùng cho PDF Optimize khác")
            return _operation_response(_operation_public(tuple(existing[:-1])))

        operation_id = str(uuid.uuid4())
        now = utc_now()
        conn.execute(
            """INSERT INTO web_document_operations
               (id, account_id, source_asset_id, project_id, kind, state, idempotency_key,
                request_fingerprint, source_sha256, source_byte_size, source_count, requested_page_range,
                created_at, queued_at, started_at, updated_at)
               VALUES (?, ?, ?, ?, ?, 'queued', ?, ?, ?, ?, 1, '', ?, ?, ?, ?)""",
            (
                operation_id, account_id, source_asset_id, str(source_row[1]) if source_row[1] else None,
                PDF_OPTIMIZE_KIND, payload.idempotency_key, request_fingerprint, source_sha256, source_bytes,
                now, now, now, now,
            ),
        )
        _record_event(conn, operation_id=operation_id, state="queued", when=now)
        conn.execute(
            "UPDATE web_document_operations SET state='processing', updated_at=? WHERE id=? AND account_id=?",
            (now, operation_id, account_id),
        )
        _record_event(conn, operation_id=operation_id, state="processing", when=now)

    try:
        source_path = _asset_path(asset_vault_directory(), source_storage_key)
        source_copy = _staging_path(root, ".source.pdf")
        _copy_verified_source(source_path, source_copy, expected_bytes=source_bytes, expected_digest=source_sha256)
        final_path, output_storage_key, output_bytes, output_digest, source_page_count = await run_in_threadpool(
            _build_optimize_output,
            root,
            source_copy,
            source_bytes=source_bytes,
        )
        saved_bytes = source_bytes - output_bytes
        now = utc_now()
        with transaction() as conn:
            current = conn.execute(
                "SELECT state FROM web_document_operations WHERE id=? AND account_id=? AND kind=?",
                (operation_id, account_id, PDF_OPTIMIZE_KIND),
            ).fetchone()
            if not current or str(current[0]) != "processing":
                raise RuntimeError("PDF Optimize không còn ở trạng thái có thể hoàn tất")
            if not _quota_available(conn, account_id=account_id, additional_bytes=output_bytes):
                raise HTTPException(status_code=413, detail="Document Operations đã đạt quota của Web account")
            conn.execute(
                """UPDATE web_document_operations
                   SET state='completed', source_page_count=?, output_page_count=?, storage_key=?,
                       original_filename='toan-aas-optimized-pdf.pdf', content_type='application/pdf',
                       byte_size=?, sha256=?, completed_at=?, updated_at=?, failure_code=NULL
                   WHERE id=? AND account_id=?""",
                (
                    source_page_count, source_page_count, output_storage_key, output_bytes, output_digest,
                    now, now, operation_id, account_id,
                ),
            )
            _record_event(conn, operation_id=operation_id, state="completed", when=now)
            _record_audit(
                conn,
                account_id=account_id,
                canonical_user_id=None,
                action="web.document_operation.pdf_optimize",
                request_id=_request_id(request),
                target=operation_id,
                detail=f"source_pages={source_page_count};saved_bytes={saved_bytes};bytes={output_bytes}",
            )
            completed = conn.execute(
                f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
                (operation_id, account_id),
            ).fetchone()
        if not completed:
            raise RuntimeError("Không thể đọc PDF Optimize vừa hoàn tất")
        final_path = None
        return _operation_response(_operation_public(tuple(completed)))
    except DocumentOperationError as exc:
        _safe_unlink(final_path)
        if exc.code == "PDF_NOT_REDUCED":
            _mark_guarded(operation_id, account_id, kind=PDF_OPTIMIZE_KIND, request=request, code=exc.code)
        else:
            if exc.code == "PDF_SOURCE_UNAVAILABLE":
                _mark_source_unavailable(source_asset_id, account_id)
            _mark_failed(operation_id, account_id, kind=PDF_OPTIMIZE_KIND, request=request, code=exc.code)
        raise HTTPException(status_code=422, detail=exc.public_message) from exc
    except HTTPException as exc:
        _safe_unlink(final_path)
        _mark_failed(
            operation_id,
            account_id,
            kind=PDF_OPTIMIZE_KIND,
            request=request,
            code="DOCUMENT_QUOTA" if exc.status_code == 413 else "DOCUMENT_OPERATION",
        )
        raise
    except Exception as exc:
        _safe_unlink(final_path)
        _mark_failed(operation_id, account_id, kind=PDF_OPTIMIZE_KIND, request=request, code="DOCUMENT_OPERATION")
        raise HTTPException(status_code=500, detail="Không thể tối ưu PDF an toàn") from exc
    finally:
        _safe_unlink(source_copy)


@router.get("/{operation_id}/download")
async def download_document_operation(operation_id: str, account: dict = Depends(require_account)):
    _require_enabled()
    operation_id = _uuid(operation_id, label="Mã thao tác tài liệu")
    account_id = str(account["id"])
    ensure_copyfast_schema()
    with transaction() as conn:
        row = conn.execute(
            f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
            (operation_id, account_id),
        ).fetchone()
    if not row or str(row[4]) != "completed":
        return _operation_not_found()
    try:
        suffix, media_type, download_filename = _output_spec(
            str(row[3] or ""),
            output_page_count=int(row[9]) if row[9] is not None else None,
        )
        # Content-type and extension are canonical server data for the known
        # operation kind. Never let a mutable database filename/MIME value
        # change how a private attachment is delivered to the browser.
        if str(row[11] or "") != media_type:
            raise RuntimeError("Artifact Document Operation có MIME không hợp lệ")
        private_path = _output_path(
            document_operations_directory(),
            str(row[19] or ""),
            expected_suffix=suffix,
        )
    except RuntimeError:
        _mark_output_unavailable(operation_id, account_id)
        return _operation_unavailable()
    if not _verify_file(private_path, expected_bytes=int(row[12] or 0), expected_digest=str(row[20] or "")):
        _mark_output_unavailable(operation_id, account_id)
        return _operation_unavailable()
    return FileResponse(
        path=private_path,
        media_type=media_type,
        filename=download_filename,
        content_disposition_type="attachment",
        headers={
            "Cache-Control": "no-store, private",
            "X-Content-Type-Options": "nosniff",
            "Referrer-Policy": "no-referrer",
            "Content-Security-Policy": "sandbox",
        },
    )


@router.get("/{operation_id}")
async def get_document_operation(operation_id: str, account: dict = Depends(require_account)):
    _require_enabled()
    operation_id = _uuid(operation_id, label="Mã thao tác tài liệu")
    ensure_copyfast_schema()
    with transaction() as conn:
        row = conn.execute(
            f"SELECT {OPERATION_SELECT} FROM web_document_operations WHERE id=? AND account_id=?",
            (operation_id, str(account["id"])),
        ).fetchone()
        events = conn.execute(
            """SELECT state, created_at FROM web_document_operation_events
               WHERE operation_id=? ORDER BY sequence ASC, id ASC LIMIT 20""",
            (operation_id,),
        ).fetchall()
    if not row:
        return _operation_not_found()
    operation = _operation_public(tuple(row))
    # Failure codes are internal branch metadata. The signed owner may see
    # lifecycle/status, but never a parser/runtime implementation hint.
    operation.pop("_failure_code", None)
    return envelope(
        True,
        "Đã tải trạng thái thao tác tài liệu.",
        data={
            "operation": operation,
            "events": [{"state": str(event[0]), "created_at": str(event[1])} for event in events],
        },
        status_name=str(row[4]) if str(row[4]) in OPERATION_STATES else "guarded",
    )
